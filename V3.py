import math
import numpy as np
import pandas as pd
import streamlit as st
from scipy.optimize import minimize
import plotly.express as px
import plotly.graph_objects as go
from typing import Dict, Any
import io
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.compose import ColumnTransformer
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.ensemble import RandomForestRegressor
from sklearn.pipeline import Pipeline

# For PDF and Word export
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from docx import Document
    EXPORT_AVAILABLE = True
except ImportError:
    EXPORT_AVAILABLE = False

# -------------------------------------------------------------------
# PAGE CONFIG
# -------------------------------------------------------------------
st.set_page_config(
    page_title="Legal Pricing Intelligence",
    page_icon="L",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# -------------------------------------------------------------------
# CSS
# -------------------------------------------------------------------
st.markdown("""
<style>
    :root {
        --primary: #3b82f6;
        --success: #10b981;
        --warning: #f59e0b;
        --purple: #8b5cf6;
    }
    .stApp { font-family: 'Segoe UI', system-ui, sans-serif; }
    #MainMenu, footer, header {visibility: hidden;}
    
    .main-header {
        text-align: center;
        padding: 1.5rem 1rem;
        margin-bottom: 1rem;
        background: linear-gradient(135deg, rgba(59,130,246,0.1) 0%, rgba(139,92,246,0.1) 100%);
        border-radius: 12px;
        border: 1px solid rgba(59,130,246,0.2);
    }
    .main-header h1 {
        font-size: 1.5rem;
        font-weight: 700;
        margin: 0;
        background: linear-gradient(135deg, #3b82f6, #8b5cf6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .main-header p { margin-top: 0.25rem; opacity: 0.7; font-size: 0.9rem; }
    
    .landing-cards-wrapper {
        display: flex;
        justify-content: center;
        gap: 2rem;
        padding: 2rem 0;
    }
    .landing-card {
        width: 320px;
        height: 250px;
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        border: 2px solid rgba(128,128,128,0.2);
        background: rgba(128,128,128,0.03);
        transition: all 0.3s ease;
    }
    .landing-card:hover {
        transform: translateY(-5px);
        border-color: var(--primary);
        box-shadow: 0 10px 30px rgba(59,130,246,0.15);
    }
    .landing-icon {
        font-size: 2rem;
        font-weight: 700;
        color: var(--primary);
        width: 60px;
        height: 60px;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-bottom: 1rem;
        border-radius: 12px;
        background: rgba(59,130,246,0.1);
    }
    .landing-title { font-size: 1.1rem; font-weight: 600; margin-bottom: 0.5rem; }
    .landing-desc { font-size: 0.85rem; opacity: 0.7; line-height: 1.5; }
    
    .panel-wrapper {
        background: rgba(128,128,128,0.03);
        border: 1px solid rgba(128,128,128,0.15);
        border-radius: 12px;
        overflow: hidden;
        margin-bottom: 0.5rem;
    }
    .panel-header {
        background: linear-gradient(135deg, rgba(59,130,246,0.1) 0%, rgba(139,92,246,0.05) 100%);
        padding: 0.75rem 1rem;
        border-bottom: 1px solid rgba(128,128,128,0.15);
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    .panel-header-icon {
        width: 28px;
        height: 28px;
        background: var(--primary);
        border-radius: 6px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-weight: 700;
        font-size: 0.8rem;
    }
    .panel-header-title { font-size: 0.95rem; font-weight: 600; margin: 0; }
    .panel-header-subtitle { font-size: 0.7rem; opacity: 0.6; margin: 0; }
    
    .panel-section {
        background: rgba(128,128,128,0.05);
        border: 1px solid rgba(128,128,128,0.1);
        border-radius: 8px;
        padding: 0.75rem;
        margin-bottom: 0.75rem;
    }
    .section-title {
        font-size: 0.75rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        color: var(--primary);
        margin-bottom: 0.5rem;
        padding-bottom: 0.25rem;
        border-bottom: 1px solid rgba(59,130,246,0.2);
    }
    
    .metric-row { display: flex; gap: 0.5rem; margin-top: 0.5rem; }
    .metric-item {
        flex: 1;
        background: rgba(128,128,128,0.05);
        border: 1px solid rgba(128,128,128,0.1);
        border-radius: 6px;
        padding: 0.5rem;
        text-align: center;
    }
    .metric-value { font-size: 1.1rem; font-weight: 700; color: var(--primary); }
    .metric-value.success { color: var(--success); }
    .metric-value.warning { color: var(--warning); }
    .metric-value.purple { color: var(--purple); }
    .metric-label { font-size: 0.65rem; text-transform: uppercase; letter-spacing: 0.3px; opacity: 0.7; }
    
    .fee-card {
        background: rgba(128,128,128,0.05);
        border: 1px solid rgba(128,128,128,0.15);
        border-radius: 8px;
        padding: 0.75rem;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .fee-card.selected {
        border-color: var(--primary);
        background: rgba(59,130,246,0.1);
    }
    .fee-card-label { font-size: 0.7rem; text-transform: uppercase; opacity: 0.7; }
    .fee-card-value { font-size: 1.25rem; font-weight: 700; margin: 0.25rem 0; }
    .fee-card-value.ml { color: #3b82f6; }
    .fee-card-value.balanced { color: #10b981; }
    .fee-card-value.risk { color: #f59e0b; }
    .fee-card-meta { font-size: 0.7rem; opacity: 0.7; }
    
    .status-badge {
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        font-weight: 600;
        font-size: 0.8rem;
        width: 100%;
        margin-top: 0.5rem;
    }
    .status-success {
        background: rgba(16,185,129,0.15);
        color: #10b981;
        border: 1px solid rgba(16,185,129,0.3);
    }
    .status-warning {
        background: rgba(245,158,11,0.15);
        color: #f59e0b;
        border: 1px solid rgba(245,158,11,0.3);
    }
    
    .phase-row {
        display: flex;
        justify-content: space-between;
        padding: 0.4rem 0;
        border-bottom: 1px solid rgba(128,128,128,0.1);
        font-size: 0.85rem;
    }
    .phase-row:last-child { border-bottom: none; }
    .phase-name { flex: 2; }
    .phase-share { flex: 1; text-align: center; opacity: 0.7; }
    .phase-amount { flex: 1; text-align: right; font-weight: 600; color: var(--primary); }
    
    .feature-pills {
        display: flex;
        flex-wrap: wrap;
        gap: 0.5rem;
        justify-content: center;
        margin-top: 1.5rem;
    }
    .feature-pill {
        padding: 0.4rem 0.8rem;
        border-radius: 20px;
        font-size: 0.75rem;
        background: rgba(59,130,246,0.1);
        border: 1px solid rgba(59,130,246,0.2);
        color: var(--primary);
    }
    
    .divider { height: 1px; background: rgba(128,128,128,0.15); margin: 0.75rem 0; }
    
    .afa-flags { display: flex; gap: 0.5rem; margin-top: 0.5rem; }
    .afa-flag { padding: 0.25rem 0.5rem; border-radius: 4px; font-size: 0.7rem; font-weight: 500; }
    .afa-flag.active { background: rgba(16,185,129,0.15); color: #10b981; border: 1px solid rgba(16,185,129,0.3); }
    .afa-flag.inactive { background: rgba(128,128,128,0.1); color: rgba(128,128,128,0.5); border: 1px solid rgba(128,128,128,0.2); }
    
    .table-header {
        display: flex;
        font-size: 0.7rem;
        font-weight: 600;
        text-transform: uppercase;
        opacity: 0.7;
        padding: 0.5rem 0;
        border-bottom: 1px solid rgba(128,128,128,0.2);
        margin-bottom: 0.5rem;
    }
    
    /* Budget Report Styling */
    .report-section {
        background: linear-gradient(135deg, rgba(59,130,246,0.05) 0%, rgba(139,92,246,0.05) 100%);
        border: 1px solid rgba(59,130,246,0.2);
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    .report-header {
        font-size: 1.2rem;
        font-weight: 700;
        color: var(--primary);
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid rgba(59,130,246,0.3);
    }
    .report-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 1rem;
        margin-bottom: 1rem;
    }
    .report-item {
        background: rgba(255,255,255,0.5);
        border: 1px solid rgba(128,128,128,0.1);
        border-radius: 8px;
        padding: 0.75rem;
    }
    .report-item-label {
        font-size: 0.7rem;
        text-transform: uppercase;
        opacity: 0.6;
        margin-bottom: 0.25rem;
    }
    .report-item-value {
        font-size: 0.95rem;
        font-weight: 600;
    }
    
    /* Professional Table Styling */
    .pro-table-container {
        background: rgba(255,255,255,0.02);
        border: 1px solid rgba(128,128,128,0.15);
        border-radius: 10px;
        overflow: hidden;
        margin-bottom: 1rem;
    }
    .pro-table-header {
        background: linear-gradient(135deg, rgba(59,130,246,0.15) 0%, rgba(139,92,246,0.1) 100%);
        padding: 0.75rem 1rem;
        font-weight: 600;
        font-size: 0.9rem;
        color: var(--primary);
        border-bottom: 1px solid rgba(59,130,246,0.2);
    }
    .pro-table-body {
        padding: 0.5rem;
    }
    .pro-table-row {
        display: flex;
        padding: 0.5rem 0.75rem;
        border-bottom: 1px solid rgba(128,128,128,0.08);
        font-size: 0.85rem;
        transition: background 0.2s;
    }
    .pro-table-row:hover {
        background: rgba(59,130,246,0.05);
    }
    .pro-table-row:last-child {
        border-bottom: none;
    }
    .pro-table-row.total {
        background: rgba(59,130,246,0.1);
        font-weight: 600;
        border-top: 2px solid rgba(59,130,246,0.3);
        margin-top: 0.5rem;
    }
    .pro-table-cell {
        flex: 1;
        text-align: center;
    }
    .pro-table-cell:first-child {
        text-align: left;
        flex: 1.5;
    }
    .pro-table-cell:last-child {
        text-align: right;
    }
    .pro-table-cell.highlight {
        color: var(--primary);
        font-weight: 600;
    }
    
    /* Summary Card */
    .summary-card {
        background: linear-gradient(135deg, rgba(16,185,129,0.1) 0%, rgba(59,130,246,0.1) 100%);
        border: 1px solid rgba(16,185,129,0.3);
        border-radius: 10px;
        padding: 1rem;
        text-align: center;
    }
    .summary-card-title {
        font-size: 0.75rem;
        text-transform: uppercase;
        opacity: 0.7;
        margin-bottom: 0.5rem;
    }
    .summary-card-value {
        font-size: 1.5rem;
        font-weight: 700;
        color: var(--success);
    }
    .summary-card-subtitle {
        font-size: 0.75rem;
        opacity: 0.6;
        margin-top: 0.25rem;
    }
    
    /* Tables side by side container */
    .tables-container {
        display: flex;
        gap: 1.5rem;
        margin: 1rem 0;
    }
    .table-half {
        flex: 1;
        min-width: 0;
    }
</style>
""", unsafe_allow_html=True)
# -------------------------------------------------------------------
# Initialize form_data in session state
# -------------------------------------------------------------------
if 'form_data' not in st.session_state:
    st.session_state.form_data = {}

# -------------------------------------------------------------------
# CONFIG
# -------------------------------------------------------------------
DATA_PATH = "ip_pricing_synthetic_10000_clean.csv"

MATTER_CATEGORIES = ["Litigation", "Corporate/M&A", "Intellectual Property", "Employment", "Real Estate", "Regulatory", "Tax", "Other"]
TECH_DOMAINS = ["AI/ML", "Biotech", "Semiconductors", "Pharma", "MedTech", "Software", "Hardware", "Other"]
CLIENT_TYPES = ["New Client", "Existing - Strategic", "Existing - Non-Strategic"]
CLIENT_INDUSTRIES = ["Technology", "Life Sciences", "Manufacturing", "Financial Services", "Energy", "Healthcare", "Other"]

IP_TYPES = {
    "Patent Drafting": "IP - Patent Drafting",
    "Prosecution": "IP - Prosecution",
    "PTAB": "IP - PTAB",
    "Litigation": "IP - Litigation",
    "FTO Analysis": "IP - FTO"
}

FEE_STRUCTURES = {
    "IP - Patent Drafting": ["Hourly", "Fixed Fee", "Tiered", "Subscription"],
    "IP - Prosecution": ["Hourly", "Fixed per OA", "Portfolio Sub", "Volume Discount"],
    "IP - PTAB": ["Phase-Based", "Fixed Fee", "Caps+Collars", "Blended"],
    "IP - Litigation": ["Blended", "Phase Caps", "Collar", "Contingency", "Success Bonus"],
    "IP - FTO": ["Fixed Fee", "Tiered", "Subscription", "Risk-Based"],
    "default": ["Hourly", "Fixed Fee", "Phase-Based", "Capped", "Blended", "Retainer"]
}

IP_PHASES = {
    "IP - Patent Drafting": [
        {"Phase": "Disclosure Review", "Weight": 0.10},
        {"Phase": "Prior Art Search", "Weight": 0.21},
        {"Phase": "Claims & Specification", "Weight": 0.52},
        {"Phase": "Drawings & Filing", "Weight": 0.17},
    ],
    "IP - Prosecution": [
        {"Phase": "Initial Filing", "Weight": 0.07},
        {"Phase": "First OA Response", "Weight": 0.17},
        {"Phase": "Second OA/RCE", "Weight": 0.26},
        {"Phase": "Examiner Interview", "Weight": 0.06},
        {"Phase": "Appeal", "Weight": 0.44},
    ],
    "IP - PTAB": [
        {"Phase": "Petition Drafting", "Weight": 0.33},
        {"Phase": "Institution Phase", "Weight": 0.11},
        {"Phase": "Trial Phase", "Weight": 0.56},
    ],
    "IP - Litigation": [
        {"Phase": "Pleadings", "Weight": 0.08},
        {"Phase": "Markman Hearing", "Weight": 0.17},
        {"Phase": "Discovery", "Weight": 0.28},
        {"Phase": "Summary Judgment", "Weight": 0.11},
        {"Phase": "Trial", "Weight": 0.36},
    ],
    "IP - FTO": [
        {"Phase": "Scoping & Mapping", "Weight": 0.10},
        {"Phase": "Search & Screening", "Weight": 0.32},
        {"Phase": "Detailed Analysis", "Weight": 0.38},
        {"Phase": "Opinion Drafting", "Weight": 0.20},
    ],
}

GENERIC_PHASES = [
    {"Phase": "Scoping & Kick-off", "Weight": 0.15},
    {"Phase": "Research & Analysis", "Weight": 0.25},
    {"Phase": "Core Execution", "Weight": 0.35},
    {"Phase": "Review & QC", "Weight": 0.15},
    {"Phase": "Delivery", "Weight": 0.10},
]

# -------------------------------------------------------------------
# ML MODEL
# -------------------------------------------------------------------
# @st.cache_resource(show_spinner=False)
# def load_model(path):
#     try:
#         df = pd.read_csv(path)
#         cat_cols = ["practice_area", "fee_structure", "client_type", "client_industry"]
#         num_cols = ["complexity_score", "duration_days", "estimated_cost", "total_hours",
#                     "client_size_revenue", "competitive_position", "competitor_avg_quote",
#                     "complexity_urgency", "team_effort", "team_size", "years_as_client",
#                     "previous_matters", "client_satisfaction", "is_subscription_model",
#                     "is_phase_based_model", "has_success_fee_component"]
#         X = df[cat_cols + num_cols]
#         y = df["target_fee"]
#         model = Pipeline([
#             ("pre", ColumnTransformer([
#                 ("cat", OneHotEncoder(handle_unknown="ignore"), cat_cols),
#                 ("num", StandardScaler(), num_cols)
#             ])),
#             ("reg", RandomForestRegressor(n_estimators=200, random_state=42, n_jobs=-1))
#         ])
#         model.fit(X, y)
#         return model
#     except Exception:
#         return None

import pickle  
import pickle

@st.cache_resource(show_spinner=False)
def load_model(path):
    try:
        with open(path, "rb") as f:
            model = pickle.load(f)
        return model
    except Exception:
        return None

model = load_model("pricing_model.pkl")



# -------------------------------------------------------------------
# OPTIMIZER
# -------------------------------------------------------------------
class FeeOptimizer:
    """Advanced Fee Optimization Engine with Dynamic Constraints"""
    
    def __init__(self, min_margin=0.15, max_margin=0.45, min_win_prob=0.30,
                 target_win_prob=0.60, cost_uncertainty=0.20, comp_uncertainty=0.10,
                 risk_tolerance=0.20):
        self.constraints = {
            'min_profit_margin': min_margin,
            'max_profit_margin': max_margin,
            'min_win_probability': min_win_prob,
            'target_win_probability': target_win_prob,
            'competitive_threshold': 0.95,
            'risk_tolerance': risk_tolerance,
            'cost_uncertainty_pct': cost_uncertainty,
            'competitor_uncertainty_pct': comp_uncertainty
        }
    
    def update_constraints(self, **kwargs):
        """Update constraints dynamically from UI"""
        for key, value in kwargs.items():
            if key in self.constraints:
                self.constraints[key] = value
    
    def calculate_win_probability(self, proposed_fee, competitor_avg, competitive_position, satisfaction=7):
        """Calculate probability of winning based on pricing and competitive factors"""
        if competitor_avg <= 0:
            return 0.5
        
        # Normalize competitive position to 0-1 range if it's in -2 to 2 range
        if competitive_position < 0:
            comp_pos_normalized = (competitive_position + 2) / 4.0
        else:
            comp_pos_normalized = np.clip(competitive_position, 0.1, 1.0)
        
        ratio = proposed_fee / competitor_avg
        
        # Sigmoid-based win probability
        base = 1.0 / (1.0 + math.exp(3.0 * (ratio - 1.0)))
        
        # Adjustments for competitive position and satisfaction
        position_bonus = 0.10 * comp_pos_normalized
        satisfaction_bonus = 0.02 * (satisfaction - 5)
        
        win_prob = base + position_bonus + satisfaction_bonus
        
        return float(np.clip(win_prob, 0.01, 0.99))
    
    def calculate_profit(self, proposed_fee, estimated_cost):
        """Calculate profit and margin"""
        if proposed_fee <= 0:
            return 0, 0
        profit = proposed_fee - estimated_cost
        margin = profit / proposed_fee
        return profit, margin
    
    def calculate_client_satisfaction_impact(self, proposed_fee, predicted_fee, current_satisfaction):
        """Estimate impact on client satisfaction"""
        if predicted_fee <= 0:
            return current_satisfaction
        
        fee_ratio = proposed_fee / predicted_fee
        
        if fee_ratio < 0.95:
            satisfaction_boost = min((0.95 - fee_ratio) * 15, 1.5)
        elif fee_ratio > 1.05:
            satisfaction_penalty = min((fee_ratio - 1.05) * 20, 2.0)
            satisfaction_boost = -satisfaction_penalty
        else:
            satisfaction_boost = 0
        
        new_satisfaction = np.clip(current_satisfaction + satisfaction_boost, 1.0, 10.0)
        return new_satisfaction
    
    def calculate_optimization_bounds(self, estimated_cost, competitor_avg, ml_prediction, strategy_type='default'):
        """Calculate valid optimization bounds for different strategies"""
        min_margin = self.constraints['min_profit_margin']
        min_fee = estimated_cost * (1 + min_margin)
        
        if strategy_type == 'aggressive':
            max_fee = max(competitor_avg * 1.1, min_fee * 1.2)
        elif strategy_type == 'conservative':
            max_fee = max(competitor_avg * 1.5, ml_prediction * 1.3, min_fee * 2.0)
        elif strategy_type == 'risk_adjusted':
            # For risk-adjusted, use uncertainty-adjusted bounds
            cost_unc = self.constraints['cost_uncertainty_pct']
            comp_unc = self.constraints['competitor_uncertainty_pct']
            worst_cost = estimated_cost * (1 + cost_unc)
            min_fee = worst_cost * (1 + min_margin)
            max_fee = max((competitor_avg + competitor_avg * comp_unc) * 1.25, min_fee * 1.3)
        else:  # default / balanced
            max_fee = max(competitor_avg * 1.3, ml_prediction * 1.2, min_fee * 1.5)
        
        if min_fee >= max_fee:
            max_fee = min_fee * 1.5
        
        return min_fee, max_fee
    
    def optimize_balanced(self, cost, competitor_avg, comp_pos, satisfaction, ml_fee, years_client=0):
        """
        Balance profit margin and win probability - optimized for expected profit.
        """
        if cost <= 0:
            return None
        
        # Ensure competitor_avg is reasonable
        if competitor_avg <= 0:
            competitor_avg = cost * 1.5
        competitor_avg = max(competitor_avg, cost * 1.2)
        
        # Normalize competitive position
        if comp_pos < 0:
            competitive_pos = (comp_pos + 2) / 4.0
        else:
            competitive_pos = np.clip(comp_pos, 0.1, 1.0)
        
        # Adjust weights based on client relationship
        if years_client > 5:
            w_profit = 0.55
            w_win_prob = 0.45
        else:
            w_profit = 0.45
            w_win_prob = 0.55
        
        # Get constraints from instance
        min_margin = self.constraints['min_profit_margin']
        max_margin = self.constraints['max_profit_margin']
        min_win_prob = self.constraints['min_win_probability']
        
        def objective(fee):
            if fee[0] <= 0:
                return 1e10
            
            profit, margin = self.calculate_profit(fee[0], cost)
            win_prob = self.calculate_win_probability(fee[0], competitor_avg, competitive_pos, satisfaction)
            
            # Penalty for constraint violations
            penalty = 0
            if margin < min_margin:
                penalty += 100 * (min_margin - margin) ** 2
            if margin > max_margin:
                penalty += 100 * (margin - max_margin) ** 2
            if win_prob < min_win_prob:
                penalty += 100 * (min_win_prob - win_prob) ** 2
            
            # Better normalization - target margin of 35% as optimal
            margin_score = np.tanh(margin * 3) if margin > 0 else 0
            
            # Nonlinear win probability scoring to encourage competitive pricing
            win_prob_score = win_prob ** 0.8
            
            # Expected profit component (profit * win_prob)
            expected_profit = profit * win_prob
            max_possible_profit = (competitor_avg * 1.2 - cost)
            expected_profit_norm = expected_profit / max_possible_profit if max_possible_profit > 0 else 0
            
            # Multi-objective: balance margin quality, win probability, and expected profit
            score = (w_profit * margin_score + 
                    w_win_prob * win_prob_score +
                    0.3 * expected_profit_norm)
            
            return -score + penalty
        
        def constraint_min_margin(fee):
            profit, margin = self.calculate_profit(fee[0], cost)
            return margin - min_margin
        
        def constraint_max_margin(fee):
            profit, margin = self.calculate_profit(fee[0], cost)
            return max_margin - margin
        
        def constraint_min_win_prob(fee):
            win_prob = self.calculate_win_probability(fee[0], competitor_avg, competitive_pos, satisfaction)
            return win_prob - min_win_prob
        
        constraints = [
            {'type': 'ineq', 'fun': constraint_min_margin},
            {'type': 'ineq', 'fun': constraint_max_margin},
            {'type': 'ineq', 'fun': constraint_min_win_prob}
        ]
        
        min_fee, max_fee = self.calculate_optimization_bounds(
            cost, competitor_avg, ml_fee, 'default'
        )
        bounds = [(min_fee, max_fee)]
        
        # Better initial guess: try multiple starting points
        initial_guesses = [
            np.clip(ml_fee, min_fee, max_fee),
            np.clip(competitor_avg * 0.95, min_fee, max_fee),
            np.clip((min_fee + max_fee) / 2, min_fee, max_fee),
            np.clip(cost * 1.30, min_fee, max_fee)
        ]
        
        best_result = None
        best_score = float('inf')
        
        for initial_guess in initial_guesses:
            try:
                result = minimize(
                    objective, 
                    [initial_guess], 
                    method='SLSQP',
                    bounds=bounds, 
                    constraints=constraints,
                    options={'maxiter': 1000, 'ftol': 1e-9}
                )
                
                if result.fun < best_score:
                    best_score = result.fun
                    best_result = result
            except Exception:
                continue
        
        if best_result:
            optimized_fee = best_result.x[0]
            profit, margin = self.calculate_profit(optimized_fee, cost)
            win_prob = self.calculate_win_probability(optimized_fee, competitor_avg, competitive_pos, satisfaction)
            
            constraints_met = {
                'min_margin': margin >= min_margin - 0.001,  # Small tolerance
                'max_margin': margin <= max_margin + 0.001,
                'min_win_prob': win_prob >= min_win_prob - 0.001
            }
            
            return {
                'fee': optimized_fee,
                'margin': margin,
                'win_prob': win_prob,
                'expected_profit': profit * win_prob,
                'profit': profit,
                'strategy': 'Balanced Multi-Objective',
                'weights_used': {
                    'profit': w_profit,
                    'win_prob': w_win_prob
                },
                'constraints_met': constraints_met,
                'all_constraints_satisfied': all(constraints_met.values()),
                'constraints_used': {
                    'min_margin': min_margin,
                    'max_margin': max_margin,
                    'min_win_prob': min_win_prob
                },
                'status': 'Success'
            }
        else:
            return None
    
    def optimize_risk_adjusted(self, cost, competitor_avg, comp_pos, satisfaction, baseline_fee):
        """
        Robust optimization considering uncertainty - uses balanced strategy result as baseline.
        """
        if cost <= 0:
            return None
        
        # Ensure competitor_avg is reasonable
        if competitor_avg <= 0:
            competitor_avg = cost * 1.5
        
        # Normalize competitive position
        if comp_pos < 0:
            competitive_pos = (comp_pos + 2) / 4.0
        else:
            competitive_pos = np.clip(comp_pos, 0.1, 1.0)
        
        # Get uncertainty parameters from constraints
        cost_uncertainty_pct = self.constraints['cost_uncertainty_pct']
        competitor_uncertainty_pct = self.constraints['competitor_uncertainty_pct']
        min_margin = self.constraints['min_profit_margin']
        min_win_prob = self.constraints['min_win_probability']
        
        cost_uncertainty = cost * cost_uncertainty_pct
        competitor_uncertainty = competitor_avg * competitor_uncertainty_pct
        
        # Worst-case scenarios
        worst_cost = cost + cost_uncertainty
        worst_competitor = max(competitor_avg - competitor_uncertainty, cost * 1.1)
        
        def objective(fee):
            if fee[0] <= 0:
                return 1e10
            
            profit = fee[0] - worst_cost
            win_prob = self.calculate_win_probability(fee[0], worst_competitor, competitive_pos, satisfaction)
            worst_case_profit = profit * win_prob
            
            # Penalty for constraint violations
            penalty = 0
            worst_margin = profit / fee[0] if fee[0] > 0 else 0
            if worst_margin < min_margin:
                penalty += 100 * (min_margin - worst_margin) ** 2
            if win_prob < 0.20:  # Minimum viability threshold
                penalty += 100 * (0.20 - win_prob) ** 2
            
            return -worst_case_profit + penalty
        
        def constraint_min_margin_worst_case(fee):
            profit = fee[0] - worst_cost
            margin = profit / fee[0] if fee[0] > 0 else 0
            return margin - min_margin
        
        def constraint_worst_case_viability(fee):
            worst_win_prob = self.calculate_win_probability(fee[0], worst_competitor, competitive_pos, satisfaction)
            return worst_win_prob - 0.20
        
        constraints = [
            {'type': 'ineq', 'fun': constraint_min_margin_worst_case},
            {'type': 'ineq', 'fun': constraint_worst_case_viability}
        ]
        
        min_fee, max_fee = self.calculate_optimization_bounds(
            cost, competitor_avg, baseline_fee, 'risk_adjusted'
        )
        bounds = [(min_fee, max_fee)]
        
        # Use baseline_fee from balanced strategy as primary initial guess
        initial_guesses = [
            np.clip(baseline_fee, min_fee, max_fee),
            np.clip((min_fee + max_fee) / 2, min_fee, max_fee),
            np.clip(baseline_fee * 1.1, min_fee, max_fee)
        ]
        
        best_result = None
        best_score = float('inf')
        
        for initial_guess in initial_guesses:
            try:
                result = minimize(
                    objective, 
                    [initial_guess], 
                    method='SLSQP',
                    bounds=bounds, 
                    constraints=constraints,
                    options={'maxiter': 1000, 'ftol': 1e-6}
                )
                
                if result.fun < best_score:
                    best_score = result.fun
                    best_result = result
            except Exception:
                continue
        
        if best_result:
            optimized_fee = best_result.x[0]
            
            # Calculate worst-case metrics
            worst_profit = optimized_fee - worst_cost
            worst_win_prob = self.calculate_win_probability(optimized_fee, worst_competitor, competitive_pos, satisfaction)
            worst_case_profit = worst_profit * worst_win_prob
            worst_margin = worst_profit / optimized_fee if optimized_fee > 0 else 0
            
            # Calculate base-case metrics
            profit, margin = self.calculate_profit(optimized_fee, cost)
            win_prob = self.calculate_win_probability(optimized_fee, competitor_avg, competitive_pos, satisfaction)
            
            return {
                'fee': optimized_fee,
                'margin': margin,
                'win_prob': win_prob,
                'expected_profit': profit * win_prob,
                'profit': profit,
                'strategy': 'Risk-Adjusted (Robust)',
                'baseline_fee_used': baseline_fee,
                'worst_case_profit': worst_case_profit,
                'worst_margin': worst_margin,
                'worst_win': worst_win_prob,
                'risk_metrics': {
                    'cost_uncertainty': cost_uncertainty,
                    'cost_uncertainty_pct': cost_uncertainty_pct,
                    'competitor_uncertainty': competitor_uncertainty,
                    'competitor_uncertainty_pct': competitor_uncertainty_pct,
                    'worst_case_cost': worst_cost,
                    'worst_case_competitor': worst_competitor,
                },
                'constraints_used': {
                    'min_margin': min_margin,
                    'min_win_prob': min_win_prob,
                    'cost_uncertainty_pct': cost_uncertainty_pct,
                    'competitor_uncertainty_pct': competitor_uncertainty_pct
                },
                'status': 'Success'
            }
        else:
            return None
    
    def win_probability(self, fee, competitor_avg, comp_pos, satisfaction=7):
        """Backward-compatible method for simple win probability calculation"""
        return self.calculate_win_probability(fee, competitor_avg, comp_pos, satisfaction)
    
    def run_all_strategies(self, cost, competitor_avg, comp_pos, satisfaction, ml_fee, years_client=0):
        """Run all optimization strategies and return results"""
        results = {}
        
        # First, run balanced strategy
        balanced_result = self.optimize_balanced(
            cost, competitor_avg, comp_pos, satisfaction, ml_fee, years_client
        )
        
        if balanced_result:
            results['balanced'] = balanced_result
            baseline_fee = balanced_result['fee']
        else:
            baseline_fee = ml_fee
        
        # Then, run risk-adjusted strategy using balanced result as baseline
        risk_result = self.optimize_risk_adjusted(
            cost, competitor_avg, comp_pos, satisfaction, baseline_fee
        )
        
        if risk_result:
            results['risk_adjusted'] = risk_result
        
        return results

# -------------------------------------------------------------------
# HELPERS
# -------------------------------------------------------------------
def compute_complexity(practice_area, **kwargs):
    if practice_area == "IP - Patent Drafting":
        base = {"Simple": 3, "Medium": 5, "Complex": 8}.get(kwargs.get('tier'), 5)
        raw = base + kwargs.get('claims', 20) / 40 + (1 if kwargs.get('drawings') else 0)
        duration = {"Simple": 30, "Medium": 45, "Complex": 60}.get(kwargs.get('tier'), 45)
    elif practice_area == "IP - Prosecution":
        raw = 4 + kwargs.get('oa_count', 2) + kwargs.get('examiner_diff', 5) / 2
        raw += 1 if kwargs.get('rce_risk') else 0
        duration = 365
    elif practice_area == "IP - PTAB":
        raw = 8 + kwargs.get('experts', 1) + kwargs.get('depositions', 4) / 5
        duration = 540
    elif practice_area == "IP - Litigation":
        raw = 6 + len(kwargs.get('phases', []))
        raw += 2 if kwargs.get('settlement_prob', 0.3) < 0.4 else 0
        duration = 730
    elif practice_area == "IP - FTO":
        scope_val = {"Narrow": 2, "Moderate": 4, "Broad": 6}.get(kwargs.get('scope'), 4)
        raw = 3 + scope_val + kwargs.get('patent_count', 50) / 100
        duration = 30
    else:
        raw = kwargs.get('general_complexity', 5)
        duration = kwargs.get('duration', 90)
    return float(np.clip(raw, 1, 10)), duration

def get_afa_flags(structure):
    s = structure or ""
    return (
        int("Subscription" in s or "Retainer" in s or "Sub" in s),
        int("Phase" in s or "Cap" in s),
        int(any(x in s for x in ["Success", "Contingency", "Risk", "Collar"]))
    )

def get_ml_prediction(model, features, total_cost):
    if model is not None:
        try:
            return float(model.predict(pd.DataFrame([features]))[0])
        except Exception:
            pass
    return total_cost * 1.35 if total_cost > 0 else 35000

# -------------------------------------------------------------------
# LANDING
# -------------------------------------------------------------------
def render_landing():
    st.markdown("""
    <div class="main-header">
        <h1>LEGAL PRICING INTELLIGENCE PLATFORM</h1>
        <p>Predictive | Optimized | Domain-Aware | Governed</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        <div class="landing-cards-wrapper">
            <div class="landing-card">
                <div class="landing-icon">M</div>
                <div class="landing-title">Matter Pricing Workbench</div>
                <div class="landing-desc">End-to-end pricing engine with ML predictions and multi-strategy optimization</div>
            </div>
            <div class="landing-card">
                <div class="landing-icon">B</div>
                <div class="landing-title">Portfolio BI Dashboard</div>
                <div class="landing-desc">Firm-level analytics, historical trends and competitive benchmarking</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Launch Workbench", key="btn_wb", use_container_width=True):
                st.session_state.page = "workbench"
                st.rerun()
        with c2:
            if st.button("Open Dashboard", key="btn_bi", use_container_width=True):
                st.session_state.page = "bi"
                st.rerun()

    st.markdown("""
    <div class="feature-pills">
        <span class="feature-pill">ML-Powered</span>
        <span class="feature-pill">Real-Time Optimization</span>
        <span class="feature-pill">Domain-Specific</span>
        <span class="feature-pill">Phase Budgeting</span>
    </div>
    """, unsafe_allow_html=True)

# -------------------------------------------------------------------
# WORKBENCH
# -------------------------------------------------------------------
def render_workbench():
    if st.button("Back to Home", key="btn_back"):
        st.session_state.page = "landing"
        st.rerun()

    st.markdown("""
    <div class="main-header">
        <h1>MATTER PRICING WORKBENCH</h1>
        <p>Real-time matter pricing with ML, optimization and governance</p>
    </div>
    """, unsafe_allow_html=True)

    model = load_model(DATA_PATH)

    col1, col2, col3 = st.columns([1, 1.2, 1])

    # ═══════════════════════════════════════════════════════════════
    # PANEL 1
    # ═══════════════════════════════════════════════════════════════
    with col1:
        st.markdown("""
        <div class="panel-wrapper">
            <div class="panel-header">
                <div class="panel-header-icon">1</div>
                <div class="panel-header-text">
                    <div class="panel-header-title">Scoping & Complexity</div>
                    <div class="panel-header-subtitle">Define matter parameters</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        with st.expander("Matter Information", expanded=True):
            client_name = st.text_input("Client Name", key="p1_client", placeholder="Enter client name...")
            matter_name = st.text_input("Matter Name", key="p1_matter", placeholder="Enter matter description...")
            c1a, c2a = st.columns(2)
            with c1a:
                matter_cat = st.selectbox("Category", MATTER_CATEGORIES, key="p1_cat")
            with c2a:
                industry = st.selectbox("Industry", CLIENT_INDUSTRIES, key="p1_ind")
        
        with st.expander("Domain Configuration", expanded=True):
            enable_ip = st.toggle("Enable IP Module", value=True, key="p1_ip_toggle")
            
            # FIXED: Initialize variables before conditional blocks
            practice_area = "Other"
            complexity_kwargs = {}
            ip_type = None  # FIXED: Initialize ip_type here
            tech_domain = None
            
            if enable_ip:
                ip_type = st.selectbox("IP Type", list(IP_TYPES.keys()), key="p1_ip_type")
                practice_area = IP_TYPES[ip_type]
                tech_domain = st.selectbox("Tech Domain", TECH_DOMAINS, key="p1_tech")
                
                st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
                
                if practice_area == "IP - Patent Drafting":
                    complexity_kwargs['tier'] = st.select_slider("Complexity", ["Simple", "Medium", "Complex"], "Medium", key="p1_tier")
                    complexity_kwargs['claims'] = st.slider("Expected Claims", 5, 100, 20, key="p1_claims")
                    complexity_kwargs['drawings'] = st.toggle("Include Drawings", True, key="p1_draw")
                elif practice_area == "IP - Prosecution":
                    complexity_kwargs['oa_count'] = st.slider("Expected OAs", 0, 5, 2, key="p1_oa")
                    complexity_kwargs['examiner_diff'] = st.slider("Examiner Difficulty", 1, 10, 6, key="p1_exam")
                    complexity_kwargs['rce_risk'] = st.toggle("RCE Risk", True, key="p1_rce")
                elif practice_area == "IP - PTAB":
                    complexity_kwargs['experts'] = st.slider("Expert Witnesses", 0, 5, 1, key="p1_exp")
                    complexity_kwargs['depositions'] = st.slider("Depositions", 0, 20, 4, key="p1_dep")
                elif practice_area == "IP - Litigation":
                    complexity_kwargs['phases'] = st.multiselect("Phases",
                        ["Pleadings", "Markman", "Discovery", "Summary Judgment", "Trial"],
                        ["Pleadings", "Markman", "Discovery"], key="p1_lit_ph")
                    complexity_kwargs['settlement_prob'] = st.slider("Settlement Probability", 0.0, 1.0, 0.3, key="p1_sett")
                elif practice_area == "IP - FTO":
                    complexity_kwargs['patent_count'] = st.slider("Patents in Scope", 1, 500, 50, key="p1_pat")
                    complexity_kwargs['scope'] = st.select_slider("Scope", ["Narrow", "Moderate", "Broad"], "Moderate", key="p1_scope")
        
        with st.expander("Complexity Drivers", expanded=True):
            if not enable_ip:
                complexity_kwargs['general_complexity'] = st.slider("Overall Complexity", 1, 10, 5, key="p1_gen")
                complexity_kwargs['duration'] = st.number_input("Duration (days)", 7, 1000, 90, step=1, key="p1_dur")
            urgency = st.slider("Urgency Level", 1, 10, 5, key="p1_urg")
            jurisdictions = st.number_input("Jurisdictions", 1, 50, 1, step=1, key="p1_jur")
            novel_issue = st.toggle("Novel Legal Issue", False, key="p1_novel")
        
        complexity_score, duration = compute_complexity(practice_area, **complexity_kwargs)
        if novel_issue:
            complexity_score = min(complexity_score + 1.5, 10)
        
        st.markdown(f"""
        <div class="panel-section">
            <div class="section-title">Derived Metrics</div>
            <div class="metric-row">
                <div class="metric-item">
                    <div class="metric-value purple">{complexity_score:.1f}</div>
                    <div class="metric-label">Complexity</div>
                </div>
                <div class="metric-item">
                    <div class="metric-value">{int(duration)}</div>
                    <div class="metric-label">Days</div>
                </div>
                <div class="metric-item">
                    <div class="metric-value">{urgency}</div>
                    <div class="metric-label">Urgency</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════
    # PANEL 2
    # ═══════════════════════════════════════════════════════════════
    with col2:
        st.markdown("""
        <div class="panel-wrapper">
            <div class="panel-header">
                <div class="panel-header-icon">2</div>
                <div class="panel-header-text">
                    <div class="panel-header-title">Team Costing & AFA</div>
                    <div class="panel-header-subtitle">Staffing, rates and fee structure</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        with st.expander("Team Staffing & Rates", expanded=True):
            st.markdown("""
            <div class="table-header">
                <div style="flex:1.5;">Role</div>
                <div style="flex:1;text-align:center;">Hours</div>
                <div style="flex:1;text-align:center;">Firm Cost/hr</div>
                <div style="flex:1;text-align:center;">Bill Rate/hr</div>
                <div style="flex:1;text-align:right;">Cost</div>
                <div style="flex:1;text-align:right;">Fee</div>
            </div>
            """, unsafe_allow_html=True)
            
            roles_config = [
                {"name": "Partner", "def_h": 12, "def_i": 300, "def_c": 950},
                {"name": "Sr. Associate", "def_h": 18, "def_i": 220, "def_c": 650},
                {"name": "Associate", "def_h": 30, "def_i": 150, "def_c": 400},
                {"name": "Paralegal", "def_h": 10, "def_i": 80, "def_c": 175}
            ]
            
            role_data = []
            for role in roles_config:
                rc1, rc2, rc3, rc4, rc5, rc6 = st.columns([1.5, 1, 1, 1, 1, 1])
                with rc1:
                    st.markdown(f"**{role['name']}**")
                with rc2:
                    hours = st.number_input("hrs", min_value=0, max_value=500, value=role['def_h'],
                                           step=1, key=f"p2h_{role['name']}", label_visibility="collapsed")
                with rc3:
                    int_rate = st.number_input("int", min_value=0, max_value=1000, value=role['def_i'],
                                              step=10, key=f"p2i_{role['name']}", label_visibility="collapsed")
                with rc4:
                    cli_rate = st.number_input("cli", min_value=0, max_value=2000, value=role['def_c'],
                                              step=10, key=f"p2c_{role['name']}", label_visibility="collapsed")
                cost = hours * int_rate
                fee = hours * cli_rate
                with rc5:
                    st.markdown(f"**${cost:,}**")
                with rc6:
                    st.markdown(f"**${fee:,}**")
                role_data.append({"Role": role['name'], "Hours": hours, "Cost": cost, "Fee": fee})
            
            total_hours = sum(r['Hours'] for r in role_data)
            total_cost = sum(r['Cost'] for r in role_data)
            total_naive = sum(r['Fee'] for r in role_data)
            
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.markdown(f"""
            <div class="metric-row">
                <div class="metric-item">
                    <div class="metric-value">{int(total_hours)}</div>
                    <div class="metric-label">Total Hours</div>
                </div>
                <div class="metric-item">
                    <div class="metric-value warning">${int(total_cost):,}</div>
                    <div class="metric-label">Firm Cost</div>
                </div>
                <div class="metric-item">
                    <div class="metric-value success">${int(total_naive):,}</div>
                    <div class="metric-label">Naive Fee</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        with st.expander("Fee Structure (AFA Model)", expanded=True):
            fee_options = FEE_STRUCTURES.get(practice_area, FEE_STRUCTURES['default'])
            fee_structure = st.selectbox("Select Model", fee_options, key="p2_fee")
            is_sub, is_phase, has_success = get_afa_flags(fee_structure)
            # st.markdown("**AFA Flags (auto-detected):**")
            # st.markdown(f"""
            # <div class="afa-flags">
            #     <span class="afa-flag {'active' if is_phase else 'inactive'}">Phase-Based</span>
            #     <span class="afa-flag {'active' if is_sub else 'inactive'}">Subscription</span>
            #     <span class="afa-flag {'active' if has_success else 'inactive'}">Success Fee</span>
            # </div>
            # """, unsafe_allow_html=True)
        
        with st.expander("Client & Market Context", expanded=True):
            c1a, c2a = st.columns(2)
            with c1a:
                client_type = st.selectbox("Client Type", CLIENT_TYPES, key="p2_ctype")
                client_size = st.number_input("Client Size ($M)", min_value=1, max_value=10000, value=900, step=10, key="p2_size")
                years_client = st.slider("Years as Client", 0, 30, 3, key="p2_yrs")
            with c2a:
                prev_matters = st.slider("Previous Matters", 0, 100, 12, key="p2_prev")
                satisfaction = st.slider("Client Satisfaction", 1, 10, 8, key="p2_sat")
                team_size = st.slider("Team Size", 1, 20, 5, key="p2_team")
            
            comp_pos = st.select_slider("Competitive Position",
                ["Very Weak", "Weak", "Neutral", "Strong", "Very Strong"], "Strong", key="p2_comp")
            comp_pos_val = {"Very Weak": -2, "Weak": -1, "Neutral": 0, "Strong": 1, "Very Strong": 2}[comp_pos]
            
            default_comp = max(int(total_naive * 0.9), int(total_cost * 1.3)) if total_cost > 0 else 28000
            competitor_avg = st.number_input("Competitor Avg Quote ($)", min_value=1000, max_value=5000000,
                                            value=default_comp, step=1000, key="p2_compavg")
        
        with st.expander("ML & Optimized Fees", expanded=True):
            features = {
                "practice_area": practice_area if practice_area.startswith("IP") else "Other / Non-IP",
                "fee_structure": fee_structure, "client_type": client_type,
                "client_industry": industry, "complexity_score": complexity_score,
                "duration_days": duration, "estimated_cost": total_cost,
                "total_hours": total_hours, "client_size_revenue": client_size,
                "competitive_position": comp_pos_val, "competitor_avg_quote": competitor_avg,
                "complexity_urgency": urgency, "team_effort": complexity_score,
                "team_size": team_size, "years_as_client": years_client,
                "previous_matters": prev_matters, "client_satisfaction": satisfaction,
                "is_subscription_model": is_sub, "is_phase_based_model": is_phase,
                "has_success_fee_component": has_success
            }
            
            ml_fee = get_ml_prediction(model, features, total_cost)
            optimizer = FeeOptimizer()
            balanced = optimizer.optimize_balanced(total_cost, competitor_avg, comp_pos_val, satisfaction, ml_fee, years_client)
            risk_adj = None
            if balanced:
                risk_adj = optimizer.optimize_risk_adjusted(total_cost, competitor_avg, comp_pos_val, satisfaction, balanced['fee'])
            
            ml_margin = ((ml_fee - total_cost) / ml_fee * 100) if ml_fee > 0 else 0
            
            st.markdown(f"""
            <div class="fee-card">
                <div class="fee-card-label">ML Baseline</div>
                <div class="fee-card-value ml">${int(ml_fee):,}</div>
                <div class="fee-card-meta">Margin: {ml_margin:.0f}%</div>
            </div>
            """, unsafe_allow_html=True)
            
            if balanced:
                st.markdown(f"""
                <div class="fee-card">
                    <div class="fee-card-label">Balanced</div>
                    <div class="fee-card-value balanced">${int(balanced['fee']):,}</div>
                    <div class="fee-card-meta">Margin: {balanced['margin']*100:.0f}% | Win: {balanced['win_prob']*100:.0f}%</div>
                </div>
                """, unsafe_allow_html=True)
            
            if risk_adj:
                st.markdown(f"""
                <div class="fee-card">
                    <div class="fee-card-label">Risk-Adjusted</div>
                    <div class="fee-card-value risk">${int(risk_adj['fee']):,}</div>
                    <div class="fee-card-meta">Margin: {risk_adj['margin']*100:.0f}% | Worst: {risk_adj['worst_margin']*100:.0f}%</div>
                </div>
                """, unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════
    # PANEL 3
    # ═══════════════════════════════════════════════════════════════
    with col3:
        st.markdown("""
        <div class="panel-wrapper">
            <div class="panel-header">
                <div class="panel-header-icon">3</div>
                <div class="panel-header-text">
                    <div class="panel-header-title">Margin & Phase Budget</div>
                    <div class="panel-header-subtitle">Final pricing and allocation</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        with st.expander("Base Fee Source", expanded=True):
            fee_options_list = ["Naive Fee", "ML Baseline"]
            if balanced:
                fee_options_list.append("Balanced")
            if risk_adj:
                fee_options_list.append("Risk-Adjusted")
            
            fee_source = st.radio("Select Base Fee", fee_options_list, horizontal=True, key="p3_source")
            
            if fee_source == "Naive Fee":
                base_fee = total_naive
            elif fee_source == "ML Baseline":
                base_fee = ml_fee
            elif fee_source == "Balanced" and balanced:
                base_fee = balanced['fee']
            elif fee_source == "Risk-Adjusted" and risk_adj:
                base_fee = risk_adj['fee']
            else:
                base_fee = ml_fee
        
        # FIXED: Optimization parameters with real-time updates
        with st.expander("Optimization Parameters", expanded=True):
            if fee_source == "Balanced":
                st.markdown("**Balanced Strategy Parameters:**")
                col_bal1, col_bal2 = st.columns(2)
                with col_bal1:
                    bal_min_margin = st.slider("Min Margin %", 5, 40, 15, key="p3_bal_min")
                    bal_max_margin = st.slider("Max Margin %", 30, 70, 50, key="p3_bal_max")
                with col_bal2:
                    bal_min_win = st.slider("Min Win Prob %", 10, 60, 25, key="p3_bal_win")
                
                # Validate max > min
                if bal_max_margin <= bal_min_margin:
                    st.warning("Max margin must be greater than min margin. Adjusting...")
                    bal_max_margin = bal_min_margin + 10
                
                # Create new optimizer with updated parameters and recalculate
                opt_balanced = FeeOptimizer(
                    min_margin=bal_min_margin/100, 
                    max_margin=bal_max_margin/100, 
                    min_win_prob=bal_min_win/100
                )
                new_balanced = opt_balanced.optimize_balanced(
                    total_cost, competitor_avg, comp_pos_val, satisfaction, ml_fee, years_client
                )
                
                if new_balanced:
                    base_fee = new_balanced['fee']
                    balanced = new_balanced  # Update balanced for later use
                    
                    # Show constraints met status
                    constraints_met = new_balanced.get('constraints_met', {})
                    all_met = new_balanced.get('all_constraints_satisfied', False)
                    
                    st.markdown(f"""
                    <div class="fee-card selected">
                        <div class="fee-card-label">Optimized Balanced Fee</div>
                        <div class="fee-card-value balanced">${int(new_balanced['fee']):,}</div>
                        <div class="fee-card-meta">Margin: {new_balanced['margin']*100:.1f}% | Win: {new_balanced['win_prob']*100:.1f}%</div>
                        <div class="fee-card-meta">Expected Profit: ${int(new_balanced['expected_profit']):,}</div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Show constraint status with details
                    st.markdown("**Constraint Status:**")
                    col_cs1, col_cs2, col_cs3 = st.columns(3)
                    with col_cs1:
                        if constraints_met.get('min_margin', False):
                            st.success(f"✓ Min Margin ({bal_min_margin}%)")
                        else:
                            st.error(f"✗ Min Margin ({bal_min_margin}%)")
                    with col_cs2:
                        if constraints_met.get('max_margin', False):
                            st.success(f"✓ Max Margin ({bal_max_margin}%)")
                        else:
                            st.error(f"✗ Max Margin ({bal_max_margin}%)")
                    with col_cs3:
                        if constraints_met.get('min_win_prob', False):
                            st.success(f"✓ Min Win ({bal_min_win}%)")
                        else:
                            st.error(f"✗ Min Win ({bal_min_win}%)")
                    
                    if all_met:
                        st.info("✓ All constraints satisfied")
                    else:
                        st.warning("⚠ Some constraints not fully met - optimizer found best feasible solution")
                else:
                    st.error("Optimization failed - using ML baseline")
                    base_fee = ml_fee
            
            elif fee_source == "Risk-Adjusted":
                st.markdown("**Risk-Adjusted Parameters:**")
                col_risk1, col_risk2 = st.columns(2)
                with col_risk1:
                    risk_cost_unc = st.slider("Cost Escalation %", 5, 50, 20, key="p3_risk_cost",
                                             help="How much costs might exceed estimates")
                with col_risk2:
                    risk_comp_unc = st.slider("Competitor Price Drop %", 5, 40, 10, key="p3_risk_comp",
                                             help="How much competitors might undercut")
                
                # Create new optimizer with updated parameters (use default min_margin and min_win_prob)
                opt_risk = FeeOptimizer(
                    cost_uncertainty=risk_cost_unc/100, 
                    comp_uncertainty=risk_comp_unc/100
                )
                
                # First get balanced result to use as baseline
                bal_for_risk = opt_risk.optimize_balanced(
                    total_cost, competitor_avg, comp_pos_val, satisfaction, ml_fee, years_client
                )
                baseline = bal_for_risk['fee'] if bal_for_risk else ml_fee
                
                new_risk = opt_risk.optimize_risk_adjusted(
                    total_cost, competitor_avg, comp_pos_val, satisfaction, baseline
                )
                
                if new_risk:
                    base_fee = new_risk['fee']
                    risk_adj = new_risk  # Update risk_adj for later use
                    
                    st.markdown(f"""
                    <div class="fee-card selected">
                        <div class="fee-card-label">Optimized Risk-Adjusted Fee</div>
                        <div class="fee-card-value risk">${int(new_risk['fee']):,}</div>
                        <div class="fee-card-meta">Margin: {new_risk['margin']*100:.1f}% | Worst Margin: {new_risk['worst_margin']*100:.1f}%</div>
                        <div class="fee-card-meta">Win Prob: {new_risk['win_prob']*100:.1f}% | Worst Win: {new_risk['worst_win']*100:.1f}%</div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Show risk metrics in a clear format
                    risk_metrics = new_risk.get('risk_metrics', {})
                    st.markdown("**Risk Scenario Analysis:**")
                    
                    col_rm1, col_rm2 = st.columns(2)
                    with col_rm1:
                        st.metric("Base Cost", f"${int(total_cost):,}")
                        st.metric("Worst-Case Cost", f"${int(risk_metrics.get('worst_case_cost', 0)):,}",
                                 delta=f"+{risk_cost_unc}%", delta_color="inverse")
                    with col_rm2:
                        st.metric("Competitor Avg", f"${int(competitor_avg):,}")
                        st.metric("Worst-Case Competitor", f"${int(risk_metrics.get('worst_case_competitor', 0)):,}",
                                 delta=f"-{risk_comp_unc}%", delta_color="inverse")
                    
                    st.markdown(f"""
                    <div class="panel-section" style="font-size: 0.85rem; margin-top: 0.5rem;">
                        <div class="section-title">Worst-Case Profit Analysis</div>
                        <div style="display: flex; justify-content: space-between; margin-top: 0.5rem;">
                            <span>Worst-Case Expected Profit:</span>
                            <span style="font-weight: 600; color: var(--warning);">${int(new_risk.get('worst_case_profit', 0)):,}</span>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.error("Risk optimization failed - using ML baseline")
                    base_fee = ml_fee
            
            elif fee_source == "ML Baseline":
                st.info(f"Using ML Baseline: ${int(ml_fee):,}")
                st.markdown("""
                <div class="panel-section">
                    <div class="section-title">About ML Baseline</div>
                    <p style="font-size: 0.85rem; opacity: 0.8;">
                        The ML baseline is a prediction from the Random Forest model trained on historical matter data. 
                        It considers practice area, fee structure, complexity, client factors, and competitive position.
                    </p>
                </div>
                """, unsafe_allow_html=True)
            
            elif fee_source == "Naive Fee":
                st.info(f"Using Naive Fee (sum of hourly rates × hours): ${int(total_naive):,}")
                naive_margin = ((total_naive - total_cost) / total_naive * 100) if total_naive > 0 else 0
                st.markdown(f"""
                <div class="panel-section">
                    <div class="section-title">Naive Fee Breakdown</div>
                    <div style="font-size: 0.85rem;">
                        <div style="display: flex; justify-content: space-between;">
                            <span>Total Hours:</span>
                            <span style="font-weight: 600;">{int(total_hours)}</span>
                        </div>
                        <div style="display: flex; justify-content: space-between;">
                            <span>Firm Cost:</span>
                            <span style="font-weight: 600;">${int(total_cost):,}</span>
                        </div>
                        <div style="display: flex; justify-content: space-between;">
                            <span>Implied Margin:</span>
                            <span style="font-weight: 600;">{naive_margin:.1f}%</span>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            else:
                st.info(f"Selected: {fee_source} - ${int(base_fee):,}")
        
        with st.expander("Fee Adjustments", expanded=True):
            required_margin = st.slider("Required Margin %", 10, 70, 35, key="p3_req")
            c1a, c2a = st.columns(2)
            with c1a:
                discount = st.slider("Discount %", 0, 40, 0, key="p3_disc")
            with c2a:
                uplift = st.slider("Uplift %", 0, 40, 0, key="p3_upl")
            
            if discount > 0 and uplift > 0:
                st.warning("Using discount only")
                uplift = 0
            
            if discount > 0:
                final_fee = base_fee * (1 - discount/100)
            elif uplift > 0:
                final_fee = base_fee * (1 + uplift/100)
            else:
                final_fee = base_fee
        
        actual_margin = ((final_fee - total_cost) / final_fee * 100) if final_fee > 0 else 0
        opt = FeeOptimizer()
        win_prob = opt.win_probability(final_fee, competitor_avg, comp_pos_val, satisfaction)
        expected_profit = (final_fee - total_cost) * win_prob
        margin_class = "success" if actual_margin >= required_margin else "warning"

        with st.expander("Fee Outcome", expanded=True):
            st.markdown(f"""
                <div class="panel-section">
                    <div class="metric-row">
                        <div class="metric-item">
                            <div class="metric-value success">${int(final_fee):,}</div>
                            <div class="metric-label">Final Fee</div>
                        </div>
                        <div class="metric-item">
                            <div class="metric-value {margin_class}">{actual_margin:.0f}%</div>
                            <div class="metric-label">Margin</div>
                        </div>
                    </div>
                    <div class="metric-row">
                        <div class="metric-item">
                            <div class="metric-value">{win_prob*100:.0f}%</div>
                            <div class="metric-label">Win Prob</div>
                        </div>
                        <div class="metric-item">
                            <div class="metric-value purple">${int(expected_profit):,}</div>
                            <div class="metric-label">Exp. Profit</div>
                        </div>
                    </div>
                </div>
            """, unsafe_allow_html=True)
        
        if actual_margin >= required_margin:
            st.markdown('<div class="status-badge status-success">Above Target - Auto-Approvable</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="status-badge status-warning">Below Target - Approval Required</div>', unsafe_allow_html=True)
        
        # Get phases for practice area
        phases = IP_PHASES.get(practice_area, GENERIC_PHASES)
        
        with st.expander("Phase Allocation", expanded=True):
            st.markdown("""
            <div style="display:flex;font-size:0.75rem;font-weight:600;opacity:0.7;padding-bottom:0.5rem;border-bottom:1px solid rgba(128,128,128,0.2);">
                <div style="flex:2;">Phase</div>
                <div style="flex:1;text-align:center;">Share</div>
                <div style="flex:1;text-align:right;">Amount</div>
            </div>
            """, unsafe_allow_html=True)
            
            total_allocated = 0
            for phase in phases:
                allocated = int(final_fee * phase['Weight'])
                total_allocated += allocated
                st.markdown(f"""
                <div class="phase-row">
                    <div class="phase-name">{phase['Phase']}</div>
                    <div class="phase-share">{int(phase['Weight']*100)}%</div>
                    <div class="phase-amount">${allocated:,}</div>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown(f"""
            <div style="display:flex;padding-top:0.5rem;margin-top:0.5rem;border-top:2px solid rgba(59,130,246,0.3);font-weight:600;">
                <div style="flex:2;">Total</div>
                <div style="flex:1;text-align:center;">100%</div>
                <div style="flex:1;text-align:right;color:var(--primary);">${total_allocated:,}</div>
            </div>
            """, unsafe_allow_html=True)
        
        
        st.session_state["final_fee"] = final_fee
        #st.write("DEBUG final_fee from session:", st.session_state.get("final_fee"))

        # ---- SAFE SESSION INIT ----
        if "form_data" not in st.session_state:
            st.session_state.form_data = {}

        # Always ensure final_fee exists
        if "final_fee" not in st.session_state:
            st.session_state["final_fee"] = 0.0
        
        #st.write("DEBUG final_fee from session:", st.session_state.get("final_fee"))
        
        # Pricing Summary Table - OUTSIDE of Phase Allocation expander
        with st.expander("Pricing Summary Table", expanded=True):

            # ---- KEY PREFIX ----
            section_key = "p3_pricing"

            # Pull total client fee from Fee Outcome
            #st.write("DEBUG final_fee from session:", st.session_state.get("final_fee"))
            total_client_fee_from_role_pricing = float(st.session_state.get("final_fee", 0.0))
            st.session_state["total_client_fee_from_role_pricing"] = total_client_fee_from_role_pricing
            #st.write("DEBUG total_client_fee_from_role_pricing:", st.session_state.get("total_client_fee_from_role_pricing"))

            # Keys
            pf_key = f"{section_key}_pf"
            db_key = f"{section_key}_db"
            dc_key = f"{section_key}_dc"
            vat_key = f"{section_key}_vat"

            # Use final_fee directly as the professional fees value
            #professional_fees_value = float(final_fee) if final_fee > 0 else float(st.session_state.form_data.get(pf_key, 0.0))

            # # Determine the value to show
            # if total_client_fee_from_role_pricing > 0:
            #     #professional_fees_value = float(total_client_fee_from_role_pricing)
            #     professional_fees_value = float(st.session_state.get("total_client_fee_from_role_pricing",0.0))
            # else:
            #     professional_fees_value = float(st.session_state.form_data.get(pf_key, 0.0))

            global professional_fees_value
            professional_fees_value = float(st.session_state.get("final_fee", 0.0))
            st.session_state["professional_fees_value"] = professional_fees_value 
            #st.write("DEBUG professional_fees_value :", st.session_state.get("professional_fees_value"))

            # professional_fees = float(st.session_state.get("professional_fees_value", 0.0))
            # st.session_state["professional_fees"] = professional_fees 
            # st.write("DEBUG professional_fees :", st.session_state.get("professional_fees"))
            # st.session_state.form_data[pf_key] = float(professional_fees)


            # ---- INPUTS ----
            # professional_fees = st.number_input(
            #     "Professional Fees ($)", 
            #     min_value=0.0, 
            #     step=100.0,
            #     value=professional_fees_value,
            #     key=f"{pf_key}_input",
            #     help="This value is auto-filled from Fee Outcome if available"
            # )
            #st.session_state.form_data[pf_key] = float(professional_fees)

            professional_fees = float(st.session_state.get("professional_fees_value", 0.0))
            st.session_state["professional_fees"] = professional_fees 
            #st.write("DEBUG professional_fees :", st.session_state.get("professional_fees"))
            st.session_state.form_data[pf_key] = float(professional_fees)




            disbursements = st.number_input(
                "Disbursements ($)", 
                min_value=0.0, 
                step=50.0, 
                value=float(st.session_state.form_data.get(db_key, 0.0)),
                key=f"{db_key}_input"
            )
            st.session_state.form_data[db_key] = float(disbursements)

            discount_pricing = st.number_input(
                "Discount ($)", 
                min_value=0.0, 
                step=50.0, 
                value=float(st.session_state.form_data.get(dc_key, 0.0)),
                key=f"{dc_key}_input"
            )
            st.session_state.form_data[dc_key] = float(discount_pricing)

            # ---- SAFE VAT HANDLING ----
            vat_value = st.session_state.form_data.get(vat_key, 1)
            if not isinstance(vat_value, int) or vat_value not in range(1, 9):
                vat_value = 1

            vat_percentage = st.selectbox(
                "VAT Percentage (%)", 
                list(range(1, 9)), 
                index=vat_value - 1,
                key=f"{vat_key}_input"
            )
            st.session_state.form_data[vat_key] = vat_percentage

            # ---- CALCULATIONS ----
            vat_amount = float(professional_fees) * (vat_percentage / 100)
            final_quoted_amount = float(professional_fees) + float(disbursements) - float(discount_pricing) + vat_amount

            # Store final quoted amount back into session
            st.session_state["final_quoted_amount"] = float(final_quoted_amount)

            # ---- TABLE ----
            pricing_summary_data = {
                "Item": [
                    "Professional Fees",
                    "Disbursements",
                    "Discount",
                    f"VAT ({vat_percentage}%)",
                    "Final Quoted Amount"
                ],
                "Amount ($)": [
                    f"{professional_fees:,.2f}",
                    f"{disbursements:,.2f}",
                    f"-{discount_pricing:,.2f}",
                    f"{vat_amount:,.2f}",
                    f"{final_quoted_amount:,.2f}"
                ]
            }
            
            pricing_summary_df = pd.DataFrame(pricing_summary_data)
            st.table(pricing_summary_df)

    # =========================================================
    # BUILD BUDGET REPORT DATA
    # =========================================================
    roles_df = pd.DataFrame(role_data)

    phase_rows = []
    for p in phases:
        phase_rows.append({
            "Phase": p["Phase"],
            "Weight": p["Weight"],
            "Amount": int(final_fee * p["Weight"])
        })

    phase_budget_df = pd.DataFrame(phase_rows)

    scoping_data = {
        "client_name": client_name,
        "matter_name": matter_name,
        "matter_category": matter_cat,
        "industry": industry,
        "domain_module": "IP" if enable_ip else "General",
        "ip_type": ip_type,  # FIXED: Now properly initialized
        "complexity_score": complexity_score,
        "duration_days": duration,
        "tech_domain": tech_domain if enable_ip else None,
    }

    costing_data = {
        "fee_structure": fee_structure,
        "total_cost": total_cost,
        "total_hours": total_hours,
        "df_roles": roles_df,
        "ml_fee": ml_fee,
        "balanced": balanced,
        "risk_adj": risk_adj,
        "total_fee_naive": total_naive,
    }

    margin_data = {
        "base_source": fee_source,
        "adj_fee": final_fee,
        "actual_margin": actual_margin / 100,
        "required_margin": required_margin / 100,
        "phase_budget_df": phase_budget_df,
        "discount_pct": discount,
        "uplift_pct": uplift,
        "win_prob": win_prob,
        "expected_profit": expected_profit,
        "pricing_summary_df": pricing_summary_df,
    }

    st.session_state["budget_report"] = {
        "scoping_data": scoping_data,
        "costing_data": costing_data,
        "margin_data": margin_data
    }

    st.session_state["final_fee"] = final_fee

# =============================================================================
# BUDGET REPORT PANEL - REDESIGNED
# =============================================================================
def panel_budget_report(
    scoping_data: Dict[str, Any],
    costing_data: Dict[str, Any],
    margin_data: Dict[str, Any],
):
    """
    Professional report-style summary with side-by-side tables
    """
    st.markdown("---")
    
    # Report Header
    st.markdown("""
    <div class="report-section">
        <div class="report-header">Budget Report Snapshot</div>
    """, unsafe_allow_html=True)
    
    # Matter Overview Grid
    st.markdown("""
        <div class="report-grid">
            <div class="report-item">
                <div class="report-item-label">Client</div>
                <div class="report-item-value">{}</div>
            </div>
            <div class="report-item">
                <div class="report-item-label">Matter</div>
                <div class="report-item-value">{}</div>
            </div>
            <div class="report-item">
                <div class="report-item-label">Category</div>
                <div class="report-item-value">{}</div>
            </div>
            <div class="report-item">
                <div class="report-item-label">Industry</div>
                <div class="report-item-value">{}</div>
            </div>
            <div class="report-item">
                <div class="report-item-label">Domain Module</div>
                <div class="report-item-value">{}</div>
            </div>
            <div class="report-item">
                <div class="report-item-label">Fee Structure</div>
                <div class="report-item-value">{}</div>
            </div>
        </div>
    """.format(
        scoping_data.get('client_name') or 'N/A',
        scoping_data.get('matter_name') or 'N/A',
        scoping_data.get('matter_category', 'N/A'),
        scoping_data.get('industry', 'N/A'),
        scoping_data.get('domain_module', 'N/A') + (f" - {scoping_data.get('ip_type')}" if scoping_data.get('ip_type') else ""),
        costing_data.get('fee_structure', 'N/A')
    ), unsafe_allow_html=True)
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Metrics Row
    total_cost = costing_data.get('total_cost', 0)
    adj_fee = margin_data.get('adj_fee', 0)
    gross_margin = (1 - (total_cost / adj_fee)) * 100 if adj_fee > 0 else 0
    
    col_m1, col_m2, col_m3, col_m4, col_m5, col_m6 = st.columns(6)
    with col_m1:
        st.metric("Complexity", f"{scoping_data.get('complexity_score', 0):0.1f}")
    with col_m2:
        st.metric("Duration", f"{int(scoping_data.get('duration_days', 0))} days")
    with col_m3:
        st.metric("Total Hours", f"{costing_data.get('total_hours', 0):0.0f}")
    with col_m4:
        st.metric("Firm Cost", f"${total_cost:,.0f}")
    with col_m5:
        st.metric("Final Fee", f"${adj_fee:,.0f}")
    with col_m6:
        st.metric("Gross Margin", f"{gross_margin:0.1f}%")
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Side-by-side Tables
    col_table1, col_table2, col_table3 = st.columns([1, 1, 1])
    
    # Staffing Mix Table
    with col_table1:
        roles_df = costing_data.get("df_roles", pd.DataFrame())
        st.markdown("""
        <div class="pro-table-container">
            <div class="pro-table-header">Staffing Mix</div>
            <div class="pro-table-body">
        """, unsafe_allow_html=True)
        
        if not roles_df.empty:
            # Header row
            st.markdown("""
                <div class="pro-table-row" style="font-weight:600;opacity:0.7;font-size:0.75rem;">
                    <div class="pro-table-cell">Role</div>
                    <div class="pro-table-cell">Hours</div>
                    <div class="pro-table-cell">Cost</div>
                    <div class="pro-table-cell">Fee</div>
                </div>
            """, unsafe_allow_html=True)
            
            for _, row in roles_df.iterrows():
                st.markdown(f"""
                    <div class="pro-table-row">
                        <div class="pro-table-cell">{row['Role']}</div>
                        <div class="pro-table-cell">{row['Hours']}</div>
                        <div class="pro-table-cell">${row['Cost']:,}</div>
                        <div class="pro-table-cell highlight">${row['Fee']:,}</div>
                    </div>
                """, unsafe_allow_html=True)
            
            # Total row
            st.markdown(f"""
                <div class="pro-table-row total">
                    <div class="pro-table-cell">Total</div>
                    <div class="pro-table-cell">{roles_df['Hours'].sum()}</div>
                    <div class="pro-table-cell">${roles_df['Cost'].sum():,}</div>
                    <div class="pro-table-cell highlight">${roles_df['Fee'].sum():,}</div>
                </div>
            """, unsafe_allow_html=True)
        
        st.markdown("</div></div>", unsafe_allow_html=True)
    
    # Phase Allocation Table
    with col_table2:
        phase_df = margin_data.get("phase_budget_df", pd.DataFrame())
        st.markdown("""
        <div class="pro-table-container">
            <div class="pro-table-header">Phase Allocation</div>
            <div class="pro-table-body">
        """, unsafe_allow_html=True)
        
        if not phase_df.empty:
            # Header row
            st.markdown("""
                <div class="pro-table-row" style="font-weight:600;opacity:0.7;font-size:0.75rem;">
                    <div class="pro-table-cell">Phase</div>
                    <div class="pro-table-cell">Share</div>
                    <div class="pro-table-cell">Amount</div>
                </div>
            """, unsafe_allow_html=True)
            
            for _, row in phase_df.iterrows():
                st.markdown(f"""
                    <div class="pro-table-row">
                        <div class="pro-table-cell">{row['Phase']}</div>
                        <div class="pro-table-cell">{row['Weight']*100:.0f}%</div>
                        <div class="pro-table-cell highlight">${row['Amount']:,}</div>
                    </div>
                """, unsafe_allow_html=True)
            
            # Total row
            st.markdown(f"""
                <div class="pro-table-row total">
                    <div class="pro-table-cell">Total</div>
                    <div class="pro-table-cell">100%</div>
                    <div class="pro-table-cell highlight">${phase_df['Amount'].sum():,}</div>
                </div>
            """, unsafe_allow_html=True)
        
        st.markdown("</div></div>", unsafe_allow_html=True)
    
    # Pricing Summary Table
    with col_table3:
        pricing_df = margin_data.get("pricing_summary_df", pd.DataFrame())
        st.markdown("""
        <div class="pro-table-container">
            <div class="pro-table-header">Pricing Summary</div>
            <div class="pro-table-body">
        """, unsafe_allow_html=True)
        
        if not pricing_df.empty:
            # Header row
            st.markdown("""
                <div class="pro-table-row" style="font-weight:600;opacity:0.7;font-size:0.75rem;">
                    <div class="pro-table-cell">Item</div>
                    <div class="pro-table-cell">Amount ($)</div>
                </div>
            """, unsafe_allow_html=True)
            
            for _, row in pricing_df.iterrows():
                is_final = "Final" in str(row['Item'])
                st.markdown(f"""
                    <div class="pro-table-row {'total' if is_final else ''}">
                        <div class="pro-table-cell">{row['Item']}</div>
                        <div class="pro-table-cell highlight">{row['Amount ($)']}</div>
                    </div>
                """, unsafe_allow_html=True)
        else:
            st.markdown("<div style='padding: 1rem; text-align: center; opacity: 0.6;'>No pricing summary available</div>", unsafe_allow_html=True)
        
        st.markdown("</div></div>", unsafe_allow_html=True)
    
    # Export buttons
    st.markdown("<br>", unsafe_allow_html=True)
    
    if EXPORT_AVAILABLE:
        df_roles = costing_data.get("df_roles", pd.DataFrame())
        pricing_summary_df = margin_data.get("pricing_summary_df", pd.DataFrame())
        
        col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 2])
        
        with col_dl1:
            # ---------------- GENERATE WORD DOCUMENT ----------------
            doc = Document()
            
            # Title
            title = doc.add_heading('REPORT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            # Section 1: Project Information
            doc.add_heading('Project Information', level=1)
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_data = [
                ('Client', scoping_data.get('client_name', 'N/A')),
                ('Matter', scoping_data.get('matter_name', 'N/A')),
                ('Category', scoping_data.get('matter_category', 'N/A')),
                ('Industry', scoping_data.get('industry', 'N/A')),
                ('Domain Module', scoping_data.get('domain_module', 'N/A')),
                ('Domain Sub-Type', scoping_data.get('ip_type', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = str(value)
                # Make label bold
                info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Section 2: Financial Summary
            doc.add_heading('Financial Summary', level=1)
            financial_table = doc.add_table(rows=7, cols=2)
            financial_table.style = 'Light Grid Accent 1'
            
            financial_data = [
                ('Fee Structure', costing_data.get('fee_structure', 'N/A')),
                ('Base Fee Source', margin_data.get('base_source', 'N/A')),
                ('Adjusted Fee', f"${margin_data['adj_fee']:,.0f}"),
                ('Estimated Cost', f"${costing_data['total_cost']:,.0f}"),
                ('Total Hours', f"{costing_data['total_hours']:0.1f}"),
                ('Gross Margin', f"{gross_margin:0.1f}%"),
                ('Required Margin', f"{margin_data['required_margin']*100:0.1f}%")
            ]
            
            for i, (label, value) in enumerate(financial_data):
                financial_table.rows[i].cells[0].text = label
                financial_table.rows[i].cells[1].text = str(value)
                financial_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Section 3: Project Metrics
            doc.add_heading('Project Metrics', level=1)
            metrics_table = doc.add_table(rows=2, cols=2)
            metrics_table.style = 'Light Grid Accent 1'
            
            metrics_data = [
                ('Complexity Score', f"{scoping_data.get('complexity_score', 0):0.1f}"),
                ('Duration (days)', f"{scoping_data.get('duration_days', 0)}")
            ]
            
            for i, (label, value) in enumerate(metrics_data):
                metrics_table.rows[i].cells[0].text = label
                metrics_table.rows[i].cells[1].text = str(value)
                metrics_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Section 4: Staffing Mix
            doc.add_heading('Staffing Mix', level=1)
            
            if not df_roles.empty:
                staffing_table = doc.add_table(rows=len(df_roles)+1, cols=len(df_roles.columns))
                staffing_table.style = 'Light Grid Accent 1'
                
                # Headers
                for j, col_name in enumerate(df_roles.columns):
                    cell = staffing_table.rows[0].cells[j]
                    cell.text = str(col_name)
                    cell.paragraphs[0].runs[0].font.bold = True
                
                # Data
                for i, row in df_roles.iterrows():
                    for j, value in enumerate(row):
                        staffing_table.rows[i+1].cells[j].text = str(value)
            
            doc.add_paragraph()
            
            # Section 5: Phase Allocation
            phase_df = margin_data.get("phase_budget_df", pd.DataFrame())
            if not phase_df.empty:
                doc.add_heading('Phase Allocation', level=1)
                
                phase_table = doc.add_table(rows=len(phase_df)+1, cols=len(phase_df.columns))
                phase_table.style = 'Light Grid Accent 1'
                
                # Headers
                for j, col_name in enumerate(phase_df.columns):
                    cell = phase_table.rows[0].cells[j]
                    cell.text = str(col_name)
                    cell.paragraphs[0].runs[0].font.bold = True
                
                # Data
                for i, row in phase_df.iterrows():
                    for j, value in enumerate(row):
                        phase_table.rows[i+1].cells[j].text = str(value)
                
                doc.add_paragraph()
            
            # Section 6: Pricing Summary
            if not pricing_summary_df.empty:
                doc.add_heading('Pricing Summary', level=1)
                pricing_table = doc.add_table(rows=len(pricing_summary_df)+1, cols=len(pricing_summary_df.columns))
                pricing_table.style = 'Light Grid Accent 1'
                
                for j, col_name in enumerate(pricing_summary_df.columns):
                    cell = pricing_table.rows[0].cells[j]
                    cell.text = str(col_name)
                    cell.paragraphs[0].runs[0].font.bold = True
                
                for i, row in pricing_summary_df.iterrows():
                    for j, value in enumerate(row):
                        pricing_table.rows[i+1].cells[j].text = str(value)
            
            # Save Word document to buffer
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            st.download_button(
                label="📄 Download Word",
                data=word_buffer,
                file_name="Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        with col_dl2:
            # ---------------- GENERATE PDF (HTML to PDF) ----------------
            phase_df = margin_data.get("phase_budget_df", pd.DataFrame())
            
            html_content = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    h1 {{ color: #1f4788; text-align: center; border-bottom: 3px solid #1f4788; padding-bottom: 10px; }}
                    h2 {{ color: #1f4788; margin-top: 30px; border-bottom: 2px solid #1f4788; padding-bottom: 5px; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th {{ background-color: #1f4788; color: white; padding: 12px; text-align: left; }}
                    td {{ border: 1px solid #ddd; padding: 10px; }}
                    tr:nth-child(even) {{ background-color: #f2f2f2; }}
                    .info-table td:first-child {{ font-weight: bold; background-color: #e8f0fe; width: 30%; }}
                </style>
            </head>
            <body>
                <h1>REPORT</h1>
                
                <h2>Project Information</h2>
                <table class="info-table">
                    <tr><td>Client</td><td>{scoping_data.get('client_name', 'N/A')}</td></tr>
                    <tr><td>Matter</td><td>{scoping_data.get('matter_name', 'N/A')}</td></tr>
                    <tr><td>Category</td><td>{scoping_data.get('matter_category', 'N/A')}</td></tr>
                    <tr><td>Industry</td><td>{scoping_data.get('industry', 'N/A')}</td></tr>
                    <tr><td>Domain Module</td><td>{scoping_data.get('domain_module', 'N/A')}</td></tr>
                    <tr><td>Domain Sub-Type</td><td>{scoping_data.get('ip_type', 'N/A')}</td></tr>
                </table>
                
                <h2>Financial Summary</h2>
                <table class="info-table">
                    <tr><td>Fee Structure</td><td>{costing_data.get('fee_structure', 'N/A')}</td></tr>
                    <tr><td>Base Fee Source</td><td>{margin_data.get('base_source', 'N/A')}</td></tr>
                    <tr><td>Adjusted Fee</td><td>${margin_data['adj_fee']:,.0f}</td></tr>
                    <tr><td>Estimated Cost</td><td>${costing_data['total_cost']:,.0f}</td></tr>
                    <tr><td>Total Hours</td><td>{costing_data['total_hours']:0.1f}</td></tr>
                    <tr><td>Gross Margin</td><td>{gross_margin:0.1f}%</td></tr>
                    <tr><td>Required Margin</td><td>{margin_data['required_margin']*100:0.1f}%</td></tr>
                </table>
                
                <h2>Project Metrics</h2>
                <table class="info-table">
                    <tr><td>Complexity Score</td><td>{scoping_data.get('complexity_score', 0):0.1f}</td></tr>
                    <tr><td>Duration (days)</td><td>{scoping_data.get('duration_days', 0)}</td></tr>
                </table>
                
                <h2>Staffing Mix</h2>
                {df_roles.to_html(index=False, border=0) if not df_roles.empty else '<p>No staffing data available</p>'}
            """
            
            # Add Phase Allocation if exists
            if not phase_df.empty:
                html_content += f"""
                <h2>Phase Allocation</h2>
                {phase_df.to_html(index=False, border=0)}
                """
            
            # Add Pricing Summary if exists
            if not pricing_summary_df.empty:
                html_content += f"""
                <h2>Pricing Summary</h2>
                {pricing_summary_df.to_html(index=False, border=0)}
                """
            
            html_content += """
            </body>
            </html>
            """
            
            # Convert HTML to PDF buffer
            pdf_buffer = BytesIO()
            pdf_buffer.write(html_content.encode('utf-8'))
            pdf_buffer.seek(0)

            st.download_button(
                label="📑 Download PDF",
                data=pdf_buffer,
                file_name="Report.html",
                mime="text/html",
                use_container_width=True
            )

# =============================================================================
# CONVERSATIONAL BI-BPC PANEL
# =============================================================================
def panel_conversational_bibpc(
    scoping_data: Dict[str, Any],
    costing_data: Dict[str, Any],
    margin_data: Dict[str, Any],
    historical_df: pd.DataFrame = None,
):
    """
    Tabbed section for explanations, what-if scenarios, and portfolio BI.
    """
    st.markdown("---")
    st.subheader("Conversational BI-BPC")

    tab_single, tab_whatif = st.tabs(
        ["Single Matter Explanation", "What-if Scenario"]
    )

    # Safe access to keys with .get() and defaults
    ml_fee = costing_data.get("ml_fee", 0)
    balanced = costing_data.get("balanced", None)
    risk_adj = costing_data.get("risk_adj", None)
    total_cost = costing_data.get("total_cost", 0)
    total_fee_naive = costing_data.get("total_fee_naive", 0)

    # ----- TAB 1: Single Matter Explanation -----
    with tab_single:
        st.markdown("### Fee Explanation")
        
        focus = st.selectbox(
            "Explanation focus",
            [
                "Why Balanced fee differs from ML baseline",
                "Why Risk-Adjusted fee differs from ML baseline",
                "Explain phase allocation",
            ],
            key="bibpc_focus"
        )

        if focus == "Why Balanced fee differs from ML baseline":
            if balanced:
                delta = balanced["fee"] - ml_fee
                direction = "higher" if delta > 0 else "lower"
                
                st.markdown(f"""
                <div class="report-section">
                    <div class="report-header">Balanced vs ML Baseline Analysis</div>
                    <div class="report-grid">
                        <div class="report-item">
                            <div class="report-item-label">ML Baseline Fee</div>
                            <div class="report-item-value">${ml_fee:,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Balanced Fee</div>
                            <div class="report-item-value" style="color: var(--success);">${balanced['fee']:,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Difference</div>
                            <div class="report-item-value">{direction} by ${abs(delta):,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Estimated Cost</div>
                            <div class="report-item-value">${total_cost:,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Balanced Margin</div>
                            <div class="report-item-value">{balanced['margin']*100:.1f}%</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Win Probability</div>
                            <div class="report-item-value">{balanced['win_prob']*100:.1f}%</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("""
                **How the Balanced Strategy Works:**
                
                The Balanced optimization strategy searches around the ML-suggested fee and identifies the optimal price point that maximizes:
                
                > **Expected Profit = (Fee - Cost) × Win Probability**
                
                Subject to:
                - A minimum margin floor (to protect profitability)
                - A maximum margin ceiling (to stay competitive)
                - A minimum win probability threshold (to ensure deal viability)
                
                The algorithm also factors in:
                - **Competitive position** - stronger positions allow higher fees
                - **Client satisfaction** - higher satisfaction supports premium pricing
                - **Years as client** - longer relationships shift focus toward profit protection
                """)
            else:
                st.info("Balanced fee was not computed for this matter. Ensure cost data is available.")

        elif focus == "Why Risk-Adjusted fee differs from ML baseline":
            if risk_adj:
                delta = risk_adj["fee"] - ml_fee
                direction = "higher" if delta > 0 else "lower"
                
                st.markdown(f"""
                <div class="report-section">
                    <div class="report-header">Risk-Adjusted vs ML Baseline Analysis</div>
                    <div class="report-grid">
                        <div class="report-item">
                            <div class="report-item-label">ML Baseline Fee</div>
                            <div class="report-item-value">${ml_fee:,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Risk-Adjusted Fee</div>
                            <div class="report-item-value" style="color: var(--warning);">${risk_adj['fee']:,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Difference</div>
                            <div class="report-item-value">{direction} by ${abs(delta):,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Estimated Cost</div>
                            <div class="report-item-value">${total_cost:,.0f}</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Risk-Adj Margin</div>
                            <div class="report-item-value">{risk_adj['margin']*100:.1f}%</div>
                        </div>
                        <div class="report-item">
                            <div class="report-item-label">Worst-Case Margin</div>
                            <div class="report-item-value" style="color: var(--warning);">{risk_adj['worst_margin']*100:.1f}%</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("""
                **How the Risk-Adjusted Strategy Works:**
                
                The Risk-Adjusted optimization strategy is designed for **downside protection**. It optimizes pricing assuming:
                
                - **Cost uncertainty** - actual costs may exceed estimates by a configured percentage
                - **Competitive uncertainty** - competitor quotes may be lower than expected
                
                The algorithm:
                1. Calculates worst-case cost (estimated cost + uncertainty buffer)
                2. Calculates worst-case competitor quote (competitor avg - uncertainty buffer)
                3. Finds the fee that maximizes expected profit under these conservative assumptions
                4. Enforces a margin floor even in the worst-case scenario
                
                This approach is ideal for:
                - Complex matters with uncertain scope
                - New practice areas with limited historical data
                - High-stakes engagements where margin protection is critical
                """)
            else:
                st.info("Risk-adjusted fee was not computed for this matter.")

        else:  # Explain phase allocation
            phase_df = margin_data.get("phase_budget_df")
            if phase_df is not None and not phase_df.empty:
                st.markdown("""
                <div class="report-section">
                    <div class="report-header">Phase Allocation Logic</div>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("""
                **How Phase Budgets Are Calculated:**
                
                The total fee is decomposed into phases using domain-specific templates:
                
                - **IP Matters**: Use specialized phase libraries (Patent Drafting, Prosecution, PTAB, Litigation, FTO)
                - **General Matters**: Use a generic phase split (Scoping, Research, Execution, Review, Delivery)
                
                Each phase receives a share of the total fee proportional to its base weight, which reflects:
                - Historical effort distribution across similar matters
                - Typical resource intensity for each phase
                - Risk and complexity concentration
                
                **Benefits of Phase Budgeting:**
                - Partners can track progress against phase-level budgets
                - Finance can forecast cash flows by phase
                - Clients receive transparency on effort allocation
                """)
                
                # Visualize phase allocation
                fig = px.pie(
                    phase_df, 
                    values='Amount', 
                    names='Phase',
                    title='Phase Budget Distribution',
                    color_discrete_sequence=px.colors.qualitative.Set2
                )
                fig.update_traces(textposition='inside', textinfo='percent+label')
                fig.update_layout(height=350, margin=dict(l=20, r=20, t=40, b=20))
                st.plotly_chart(fig, use_container_width=True)
                
                st.dataframe(phase_df, use_container_width=True, hide_index=True)
            else:
                st.info("No phase budget available for this matter.")

    # ----- TAB 2: What-if Scenario -----
    with tab_whatif:
        st.markdown("### What-if Scenario Analysis")
        
        st.markdown("""
        <div class="panel-section">
            <div class="section-title">Adjust Parameters to See Impact</div>
        </div>
        """, unsafe_allow_html=True)

        base_fee_source = margin_data.get("base_source", "ML Baseline")
        
        if base_fee_source == "Naive Fee":
            base_fee = total_fee_naive
        elif base_fee_source == "ML Baseline":
            base_fee = ml_fee
        elif base_fee_source == "Balanced":
            base_fee = balanced["fee"] if balanced else ml_fee
        elif base_fee_source == "Risk-Adjusted":
            base_fee = risk_adj["fee"] if risk_adj else ml_fee
        else:
            base_fee = ml_fee

        current_required = margin_data.get("required_margin", 0.35)
        current_discount = margin_data.get("discount_pct", 0)

        col_wif1, col_wif2 = st.columns(2)
        
        with col_wif1:
            new_required_margin = st.slider(
                "Target Margin (%)",
                10, 70,
                int(current_required * 100),
                1,
                key="whatif_margin",
                help="Adjust the required margin threshold"
            )
            
            new_discount = st.slider(
                "Apply Discount (%)",
                0, 40,
                int(current_discount),
                1,
                key="whatif_discount",
                help="Simulate a client discount"
            )
        
        with col_wif2:
            new_cost_change = st.slider(
                "Cost Change (%)",
                -30, 50,
                0,
                5,
                key="whatif_cost",
                help="Simulate cost overrun or savings"
            )
            
            # new_hours_change = st.slider(
            #     "Hours Change (%)",
            #     -30, 50,
            #     0,
            #     5,
            #     key="whatif_hours",
            #     help="Simulate effort variance"
            # )

        # Calculate what-if scenario
        whatif_cost = total_cost * (1 + new_cost_change / 100.0)
        
        if new_discount > 0:
            whatif_fee = base_fee * (1 - new_discount / 100.0)
        else:
            whatif_fee = base_fee

        if whatif_fee > 0:
            whatif_margin = (whatif_fee - whatif_cost) / whatif_fee
        else:
            whatif_margin = 0.0

        whatif_profit = whatif_fee - whatif_cost
        
        # Display results
        st.markdown("<br>", unsafe_allow_html=True)
        
        col_res1, col_res2, col_res3, col_res4 = st.columns(4)
        
        with col_res1:
            cost_delta = whatif_cost - total_cost
            st.metric(
                "Adjusted Cost", 
                f"${whatif_cost:,.0f}",
                delta=f"${cost_delta:+,.0f}" if cost_delta != 0 else None,
                delta_color="inverse"
            )
        
        with col_res2:
            fee_delta = whatif_fee - base_fee
            st.metric(
                "Adjusted Fee", 
                f"${whatif_fee:,.0f}",
                delta=f"${fee_delta:+,.0f}" if fee_delta != 0 else None
            )
        
        with col_res3:
            original_margin = margin_data.get("actual_margin", 0)
            margin_delta = whatif_margin - original_margin
            st.metric(
                "New Margin", 
                f"{whatif_margin*100:.1f}%",
                delta=f"{margin_delta*100:+.1f}%" if margin_delta != 0 else None
            )
        
        with col_res4:
            st.metric(
                "Gross Profit", 
                f"${whatif_profit:,.0f}"
            )
        
        # Status check
        st.markdown("<br>", unsafe_allow_html=True)
        
        if whatif_margin >= new_required_margin / 100.0:
            st.markdown("""
            <div class="status-badge status-success">
                ✓ This scenario meets the target margin requirement
            </div>
            """, unsafe_allow_html=True)
        else:
            shortfall = (new_required_margin / 100.0 - whatif_margin) * whatif_fee
            st.markdown(f"""
            <div class="status-badge status-warning">
                ⚠ This scenario falls short by ${shortfall:,.0f} to meet the {new_required_margin}% margin target
            </div>
            """, unsafe_allow_html=True)
            
            # Suggest corrective action
            required_fee = whatif_cost / (1 - new_required_margin / 100.0)
            st.info(f"💡 To achieve {new_required_margin}% margin with current costs, the fee would need to be **${required_fee:,.0f}**")

    # ----- TAB 3: Portfolio BI View -----
    # with tab_port:
    #     st.markdown("### Portfolio-Level Pricing Intelligence")

    #     if historical_df is None or historical_df.empty:
    #         st.info("Connect a historical matters dataset to enable firm-wide BI analytics.")
            
    #         st.markdown("<br>", unsafe_allow_html=True)
            
    #         # Show sample placeholder with demo data
    #         st.markdown("""
    #         <div class="report-section">
    #             <div class="report-header">Sample Portfolio Metrics (Demo Data)</div>
    #         </div>
    #         """, unsafe_allow_html=True)
            
    #         # Demo data visualization
    #         col_demo1, col_demo2 = st.columns(2)
            
    #         with col_demo1:
    #             sample_margin = pd.DataFrame({
    #                 'Practice Area': ['IP - Patent Drafting', 'IP - Prosecution', 'IP - PTAB', 'IP - Litigation', 'IP - FTO'],
    #                 'Avg Margin (%)': [32.5, 28.4, 41.2, 35.8, 29.1],
    #                 'Matters': [45, 82, 23, 31, 56]
    #             })
                
    #             fig = px.bar(
    #                 sample_margin, 
    #                 x='Practice Area', 
    #                 y='Avg Margin (%)',
    #                 color='Avg Margin (%)',
    #                 color_continuous_scale='Greens',
    #                 title='Average Margin by Practice Area'
    #             )
    #             fig.update_layout(height=300, margin=dict(l=20, r=20, t=40, b=20))
    #             st.plotly_chart(fig, use_container_width=True)
            
    #         with col_demo2:
    #             sample_win = pd.DataFrame({
    #                 'Practice Area': ['IP - Patent Drafting', 'IP - Prosecution', 'IP - PTAB', 'IP - Litigation', 'IP - FTO'],
    #                 'Win Rate (%)': [68, 72, 58, 65, 74],
    #                 'Proposals': [66, 114, 40, 48, 76]
    #             })
                
    #             fig = px.bar(
    #                 sample_win, 
    #                 x='Practice Area', 
    #                 y='Win Rate (%)',
    #                 color='Win Rate (%)',
    #                 color_continuous_scale='Blues',
    #                 title='Win Rate by Practice Area'
    #             )
    #             fig.update_layout(height=300, margin=dict(l=20, r=20, t=40, b=20))
    #             st.plotly_chart(fig, use_container_width=True)
            
    #         # Sample combined table
    #         st.markdown("**Combined Performance Summary:**")
    #         combined_sample = pd.DataFrame({
    #             'Practice Area': ['IP - Patent Drafting', 'IP - Prosecution', 'IP - PTAB', 'IP - Litigation', 'IP - FTO'],
    #             'Avg Margin (%)': [32.5, 28.4, 41.2, 35.8, 29.1],
    #             'Win Rate (%)': [68, 72, 58, 65, 74],
    #             'Matters': [45, 82, 23, 31, 56],
    #             'Total Revenue ($K)': [1250, 890, 1520, 2100, 620]
    #         })
    #         st.dataframe(combined_sample, use_container_width=True, hide_index=True)
            
    #         return

        # # Real data analysis
        # view_type = st.selectbox(
        #     "Choose BI View",
        #     [
        #         "Average margin by practice area & AFA",
        #         "Win rate by practice area & AFA",
        #         "Revenue by practice area",
        #         "Margin trend over time",
        #     ],
        #     key="portfolio_view"
        # )

        # group_cols = ["practice_area", "fee_structure"]
        # df = historical_df.copy()

        # if view_type == "Average margin by practice area & AFA":
        #     if "realized_margin_pct" not in df.columns:
        #         st.warning("Column 'realized_margin_pct' not found in dataset.")
        #         return
            
        #     df["realized_margin_pct"] = df["realized_margin_pct"].astype(float)
        #     agg = df.groupby(group_cols)["realized_margin_pct"].mean().reset_index()
        #     agg["Avg Realized Margin (%)"] = (agg["realized_margin_pct"] * 100).round(1)
            
        #     fig = px.bar(
        #         agg, 
        #         x='practice_area', 
        #         y='Avg Realized Margin (%)',
        #         color='fee_structure',
        #         barmode='group',
        #         title='Average Realized Margin by Practice Area & Fee Structure'
        #     )
        #     fig.update_layout(height=400, margin=dict(l=20, r=20, t=40, b=20))
        #     st.plotly_chart(fig, use_container_width=True)
            
        #     st.dataframe(agg.drop(columns=["realized_margin_pct"]), use_container_width=True, hide_index=True)
        #     st.caption("Shows where the firm systematically protects or loses margin across practice areas and fee models.")
            
        # elif view_type == "Win rate by practice area & AFA":
        #     if "won_matter_flag" not in df.columns:
        #         st.warning("Column 'won_matter_flag' not found in dataset.")
        #         return
            
        #     agg = df.groupby(group_cols)["won_matter_flag"].mean().reset_index()
        #     agg["Win Rate (%)"] = (agg["won_matter_flag"] * 100).round(1)
            
        #     fig = px.bar(
        #         agg, 
        #         x='practice_area', 
        #         y='Win Rate (%)',
        #         color='fee_structure',
        #         barmode='group',
        #         title='Win Rate by Practice Area & Fee Structure'
        #     )
        #     fig.update_layout(height=400, margin=dict(l=20, r=20, t=40, b=20))
        #     st.plotly_chart(fig, use_container_width=True)
            
        #     st.dataframe(agg.drop(columns=["won_matter_flag"]), use_container_width=True, hide_index=True)
        #     st.caption("Shows which combinations of practice area and fee model are commercially attractive (high win rates).")
            
        # elif view_type == "Revenue by practice area":
        #     if "target_fee" not in df.columns:
        #         st.warning("Column 'target_fee' not found in dataset.")
        #         return
            
        #     agg = df.groupby("practice_area")["target_fee"].sum().reset_index()
        #     agg["Revenue ($M)"] = (agg["target_fee"] / 1_000_000).round(2)
            
        #     fig = px.pie(
        #         agg, 
        #         values='Revenue ($M)', 
        #         names='practice_area',
        #         title='Revenue Distribution by Practice Area'
        #     )
        #     fig.update_layout(height=400, margin=dict(l=20, r=20, t=40, b=20))
        #     st.plotly_chart(fig, use_container_width=True)
            
        # else:  # Margin trend over time
        #     if "realized_margin_pct" not in df.columns or "created_date" not in df.columns:
        #         st.warning("Required columns not found for trend analysis.")
        #         return
            
        #     df["created_date"] = pd.to_datetime(df["created_date"])
        #     df["month"] = df["created_date"].dt.to_period("M").astype(str)
            
        #     trend = df.groupby("month")["realized_margin_pct"].mean().reset_index()
        #     trend["Avg Margin (%)"] = (trend["realized_margin_pct"] * 100).round(1)
            
        #     fig = px.line(
        #         trend, 
        #         x='month', 
        #         y='Avg Margin (%)',
        #         markers=True,
        #         title='Average Margin Trend Over Time'
        #     )
        #     fig.update_traces(line_color='#3b82f6', line_width=3)
        #     fig.update_layout(height=400, margin=dict(l=20, r=20, t=40, b=20))
        #     st.plotly_chart(fig, use_container_width=True)


# -------------------------------------------------------------------
# BI DASHBOARD
# -------------------------------------------------------------------
def render_bi():
    if st.button("← Back to Home", key="btn_back_bi"):
        st.session_state.page = "landing"
        st.rerun()

    st.markdown("""
    <div class="main-header">
        <h1>PORTFOLIO BI & BENCHMARKING</h1>
        <p>Firm-level analytics and competitive intelligence</p>
    </div>
    """, unsafe_allow_html=True)

    st.info("🚀 Coming Soon - Portfolio analytics, historical trends, and competitive benchmarking")
    
    st.markdown("<br>", unsafe_allow_html=True)

    # Demo visualizations
    c1, c2 = st.columns(2)
    with c1:
        data = pd.DataFrame({
            'Practice': ['Drafting', 'Prosecution', 'PTAB', 'Litigation', 'FTO'],
            'Revenue': [1.2, 0.8, 1.5, 2.1, 0.6]
        })
        fig = px.bar(
            data, 
            x='Practice', 
            y='Revenue', 
            color='Practice',
            color_discrete_sequence=px.colors.qualitative.Set2,
            title="Revenue by Practice Area ($M)"
        )
        fig.update_layout(height=350, showlegend=False, margin=dict(l=20, r=20, t=40, b=20))
        st.plotly_chart(fig, use_container_width=True)
        
    with c2:
        data = pd.DataFrame({
            'Month': ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun'],
            'Rate': [62, 58, 65, 70, 68, 72]
        })
        fig = px.line(
            data, 
            x='Month', 
            y='Rate', 
            markers=True,
            title="Win Rate Trend (%)"
        )
        fig.update_traces(line_color='#3b82f6', line_width=3, marker_size=10)
        fig.update_layout(height=350, margin=dict(l=20, r=20, t=40, b=20))
        st.plotly_chart(fig, use_container_width=True)
    
    # Additional demo charts
    c3, c4 = st.columns(2)
    with c3:
        data = pd.DataFrame({
            'Fee Structure': ['Hourly', 'Fixed Fee', 'Phase-Based', 'Blended', 'Subscription'],
            'Count': [45, 32, 28, 18, 12]
        })
        fig = px.pie(
            data, 
            values='Count', 
            names='Fee Structure',
            title="Matters by Fee Structure",
            color_discrete_sequence=px.colors.qualitative.Pastel
        )
        fig.update_layout(height=350, margin=dict(l=20, r=20, t=40, b=20))
        st.plotly_chart(fig, use_container_width=True)
        
    with c4:
        data = pd.DataFrame({
            'Client Type': ['Strategic', 'Non-Strategic', 'New Client'],
            'Margin': [38, 32, 28],
            'Win Rate': [75, 65, 55]
        })
        fig = go.Figure()
        fig.add_trace(go.Bar(name='Margin %', x=data['Client Type'], y=data['Margin'], marker_color='#10b981'))
        fig.add_trace(go.Bar(name='Win Rate %', x=data['Client Type'], y=data['Win Rate'], marker_color='#3b82f6'))
        fig.update_layout(
            barmode='group', 
            height=350, 
            title="Performance by Client Type",
            margin=dict(l=20, r=20, t=40, b=20)
        )
        st.plotly_chart(fig, use_container_width=True)


# -------------------------------------------------------------------
# MAIN
# -------------------------------------------------------------------
def main():
    if 'page' not in st.session_state:
        st.session_state.page = "landing"

    if st.session_state.page == "landing":
        render_landing()

    elif st.session_state.page == "workbench":
        render_workbench()

        # Render additional panels if budget report data exists
        if "budget_report" in st.session_state:
            br = st.session_state["budget_report"]

            # Budget report panel
            panel_budget_report(
                br["scoping_data"],
                br["costing_data"],
                br["margin_data"]
            )

            # Conversational BI-BPC panel
            panel_conversational_bibpc(
                br["scoping_data"],
                br["costing_data"],
                br["margin_data"],
                historical_df=None  # Pass actual historical_df if available
            )

    elif st.session_state.page == "bi":
        render_bi()


if __name__ == "__main__":
    main()